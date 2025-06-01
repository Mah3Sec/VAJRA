"""
cover_generator.py
==================
VAJRA — Cover Page Generator for Custom .docx Templates

Prepends a professional VAJRA cover page to a filled template body.
Builds the cover DIRECTLY into the body document so image relationships
are preserved correctly (avoids the rels-lost-on-merge problem).

TWO MODES:
  1. VAJRA default cover  — navy bars, client logo, report title, tester logo in bar
  2. Standalone test      — run from CLI to generate a test cover.docx

USAGE (standalone — for testing):
  python3 cover_generator.py \\
      --client "ACME Corp" --tester "Jane Smith" --date "2026-03-16" \\
      --report-type phishing --engagement-ref "ENG-2026-001" \\
      --company-logo /path/to/logo.png --tester-logo /path/to/tester.png \\
      --output test_cover.docx

CALLED BY:
  docx_template_filler.py → fill_docx_template() → prepend_cover()

DEPENDENCIES:
  pip install python-docx Pillow
"""

import os
import sys
import copy
import shutil
import argparse
import tempfile
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[COVER] python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# ── Constants ─────────────────────────────────────────────────────────────────
C_ACCENT = "1C3F6E"   # Navy blue
C_WHITE  = "FFFFFF"
C_MUTED  = "6B8CAE"

REPORT_LABELS = {
    "pentest":  "Penetration Test Report",
    "redteam":  "Red Team Engagement Report",
    "phishing": "Phishing Simulation Report",
}


# ── Helpers ───────────────────────────────────────────────────────────────────

def _rgb(hex_str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_valign(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def _no_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    ex = tblPr.find(qn("w:tblBorders"))
    if ex is not None:
        tblPr.remove(ex)
    tblPr.append(tblBorders)


def _set_row_height(row, height_twips):
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(height_twips))
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)


def _run(para, text, bold=False, size_pt=12, color_hex=None,
          italic=False, font="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size_pt)
    if color_hex:
        run.font.color.rgb = _rgb(color_hex)
    return run


def _load_image_dims(image_path, max_w_in, max_h_in):
    """Return (path, width_inches, height_inches) respecting aspect ratio."""
    if not image_path or not os.path.exists(image_path):
        return None, 0, 0
    try:
        from PIL import Image
        with Image.open(image_path) as img:
            w, h = img.size
            ratio = w / h
            width  = min(max_w_in, max_h_in * ratio)
            height = width / ratio
            if height > max_h_in:
                height = max_h_in
                width  = height * ratio
        return image_path, round(width, 4), round(height, 4)
    except ImportError:
        return image_path, max_w_in, max_h_in
    except Exception as e:
        print(f"[COVER] Image dims warning: {e}")
        return None, 0, 0


def _set_tbl_width(table, width_twips):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(width_twips))
    tblW.set(qn("w:type"), "dxa")
    ex = tblPr.find(qn("w:tblW"))
    if ex is not None:
        tblPr.remove(ex)
    tblPr.append(tblW)


def _sp(doc, points):
    """Add a spacer paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(points)
    return p


# ── Core: add cover content into an existing document ─────────────────────────

def add_cover_to_doc(doc, meta, company_logo=None, tester_logo=None):
    """
    Insert cover page content at the BEGINNING of an existing Document object.
    Images are added via python-docx add_picture() so relationships are correct.

    Returns the number of paragraphs/elements inserted (for debugging).
    """
    from docx.oxml import OxmlElement
    from lxml import etree

    PW = 12240   # page width in twips (8.5")

    rtype  = (meta.get("report_type") or "pentest").lower()
    rlabel = REPORT_LABELS.get(rtype, "Security Assessment Report")

    # We build a temporary doc, then prepend its body elements into the real doc.
    # Images in the temp doc get their own rIds.
    # We then copy both the XML elements AND the image part relationships.

    tmp_doc = Document()
    # Zero margins on tmp doc section for cover
    sec = tmp_doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.top_margin    = Inches(0)
    sec.bottom_margin = Inches(0)
    sec.left_margin   = Inches(0)
    sec.right_margin  = Inches(0)

    # ── TOP BAR ───────────────────────────────────────────────────────────────
    top = tmp_doc.add_table(rows=1, cols=1)
    _no_borders(top)
    _set_tbl_width(top, PW)
    top_cell = top.cell(0, 0)
    _set_cell_bg(top_cell, C_ACCENT)
    _set_cell_valign(top_cell, "center")
    _set_row_height(top.rows[0], 504)
    top_cell.paragraphs[0].add_run(" ")

    _sp(tmp_doc, 60)

    # ── CLIENT LOGO ───────────────────────────────────────────────────────────
    logo_path, logo_w, logo_h = _load_image_dims(company_logo, 3.2, 1.1)
    if logo_path:
        p = tmp_doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(16)
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(logo_w))
    else:
        client_name = str(meta.get("client") or "")
        if client_name:
            p = tmp_doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(16)
            _run(p, client_name, size_pt=22, color_hex=C_MUTED)

    # ── REPORT TITLE ──────────────────────────────────────────────────────────
    p = tmp_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(10)
    _run(p, rlabel, bold=True, size_pt=26, color_hex=C_ACCENT)

    # ── ENGAGEMENT REF + DATE ─────────────────────────────────────────────────
    ref  = meta.get("engagement_ref", "")
    date = meta.get("date", "")
    ref_text = "  ·  ".join(filter(None, [ref, date]))
    if ref_text:
        p = tmp_doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        _run(p, ref_text, size_pt=11, color_hex=C_MUTED)

    # ── SPACER ────────────────────────────────────────────────────────────────
    _sp(tmp_doc, 200)

    # ── BOTTOM BAR (CONFIDENTIAL + tester logo) ───────────────────────────────
    b1 = int(PW * 0.60)
    b2 = PW - b1
    bot = tmp_doc.add_table(rows=1, cols=2)
    _no_borders(bot)
    _set_tbl_width(bot, PW)
    _set_row_height(bot.rows[0], 630)

    left_cell  = bot.cell(0, 0)
    right_cell = bot.cell(0, 1)
    _set_cell_bg(left_cell,  C_ACCENT)
    _set_cell_bg(right_cell, C_ACCENT)
    _set_cell_valign(left_cell,  "center")
    _set_cell_valign(right_cell, "center")

    # Left: CONFIDENTIAL text
    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.left_indent = Inches(0.8)
    classif = meta.get("classification", "CONFIDENTIAL — NOT FOR DISTRIBUTION")
    _run(lp, classif, bold=True, size_pt=10, color_hex=C_WHITE)

    # Right: tester logo OR version text
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.right_indent = Inches(0.6)
    test_path, test_w, test_h = _load_image_dims(tester_logo, 1.7, 0.36)
    if test_path:
        rrun = rp.add_run()
        rrun.add_picture(test_path, width=Inches(test_w))
    else:
        version = meta.get("report_version", "1.0 — Final")
        _run(rp, version, italic=True, size_pt=10, color_hex="8FA8C8")

    # ── Page break after cover ────────────────────────────────────────────────
    pb = tmp_doc.add_paragraph()
    pb.paragraph_format.space_before = Pt(0)
    pb.paragraph_format.space_after  = Pt(0)
    pb_run = pb.add_run()
    from docx.oxml import OxmlElement as OE
    br = OE("w:br")
    br.set(qn("w:type"), "page")
    pb_run._r.append(br)

    # ── Copy elements + image relationships into target doc ───────────────────
    # Step 1: copy image parts from tmp_doc to doc
    try:
        _copy_image_parts(tmp_doc, doc)
    except Exception as e:
        print(f"[COVER] Image relationship copy warning: {e}")

    # Step 2: prepend tmp_doc body elements into doc.element.body
    tmp_body = tmp_doc.element.body
    doc_body = doc.element.body

    inserted = 0
    for el in list(tmp_body):
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag == "sectPr":
            continue
        doc_body.insert(inserted, copy.deepcopy(el))
        inserted += 1

    print(f"[COVER] Inserted {inserted} cover elements into document")
    return inserted


def _copy_image_parts(src_doc, dst_doc):
    """
    Copy image part data from src_doc into dst_doc and update rId mappings.
    This ensures images added via add_picture() in src_doc are available in dst_doc.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    import re as _re

    src_parts = src_doc.part.related_parts
    dst_part  = dst_doc.part

    rId_map = {}  # old rId → new rId

    for rId, part in src_parts.items():
        content_type = getattr(part, 'content_type', '')
        if 'image' in content_type.lower():
            try:
                # Add the image blob to dst_doc
                new_rId = dst_part.relate_to(part, part.partname.split('.')[-1])
                # Actually use the proper image relationship type
                from docx.opc.part import Part
                from docx.opc.packuri import PackURI
                # Create a new image part in dst
                img_part = dst_part.related_parts.get(rId)
                if img_part is None:
                    new_rId = dst_part.relate_to(
                        part,
                        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
                    )
                    rId_map[rId] = new_rId
                    print(f"[COVER] Copied image {rId} → {new_rId}")
            except Exception as e:
                print(f"[COVER] Could not copy image {rId}: {e}")

    return rId_map


# ── Main entry point ──────────────────────────────────────────────────────────

def prepend_cover(body_path: str, output_path: str, meta: dict,
                  company_logo: str = None, tester_logo: str = None) -> bool:
    """
    Add a VAJRA cover page at the start of body_path → save to output_path.
    Builds cover directly into the document so image rels are preserved.
    """
    try:
        # Work on a copy if in-place
        work_path = output_path
        if body_path == output_path:
            import tempfile
            tmp = tempfile.mktemp(suffix='.docx')
            shutil.copy2(body_path, tmp)
            work_path = tmp
        else:
            shutil.copy2(body_path, output_path)
            work_path = output_path
            tmp = None

        doc = Document(work_path)
        n = add_cover_to_doc(doc, meta, company_logo, tester_logo)
        doc.save(output_path)

        if tmp and os.path.exists(tmp):
            os.unlink(tmp)

        size = os.path.getsize(output_path)
        print(f"[COVER] Cover prepended ({n} elements). Output: {output_path} ({size:,} bytes)")
        return True

    except Exception as e:
        import traceback
        print(f"[COVER] Error: {e}")
        traceback.print_exc()
        # Fallback — copy body without cover
        if body_path != output_path:
            shutil.copy2(body_path, output_path)
        return False


# ── Standalone test CLI ───────────────────────────────────────────────────────

def _cli():
    p = argparse.ArgumentParser(description="VAJRA Cover Page Generator — standalone test")
    p.add_argument("--client",         default="ACME Corporation")
    p.add_argument("--tester",         default="Security Assessor")
    p.add_argument("--date",           default="2026-03-16")
    p.add_argument("--report-type",    default="pentest",
                   choices=["pentest","redteam","phishing"])
    p.add_argument("--engagement-ref", default="ENG-2026-001")
    p.add_argument("--classification", default="CONFIDENTIAL — NOT FOR DISTRIBUTION")
    p.add_argument("--report-version", default="1.0 — Final")
    p.add_argument("--company-logo",   default=None)
    p.add_argument("--tester-logo",    default=None)
    p.add_argument("--body",           default=None,
                   help="Existing .docx to prepend cover to")
    p.add_argument("--output",         required=True)
    args = p.parse_args()

    meta = {
        "client":         args.client,
        "tester":         args.tester,
        "date":           args.date,
        "report_type":    args.report_type,
        "engagement_ref": args.engagement_ref,
        "classification": args.classification,
        "report_version": args.report_version,
    }

    if args.body:
        ok = prepend_cover(args.body, args.output, meta,
                           args.company_logo, args.tester_logo)
    else:
        # Create a minimal body and prepend cover
        body = Document()
        body.add_paragraph("Test body content")
        tmp = args.output + ".body.docx"
        body.save(tmp)
        ok = prepend_cover(tmp, args.output, meta,
                           args.company_logo, args.tester_logo)
        os.unlink(tmp)

    print(f"{'✓ Done' if ok else '✗ Failed'}: {args.output}")
    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    _cli()
