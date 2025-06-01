"""
VAJRA — Vulnerability Analysis, Judgement & Reporting Arsenal
Runs 100% on localhost:5000 — no data leaves your machine except to your configured AI endpoint.
"""

import os
import sys
import json
import glob
import sqlite3
import traceback
try:
    import defusedxml.ElementTree as ET
except ImportError:
    import xml.etree.ElementTree as ET
    import warnings
    warnings.warn("[SECURITY] defusedxml not installed — XML parsing may be vulnerable to XXE. Run: pip3 install defusedxml")
import threading
import concurrent.futures
import re
from datetime import datetime, timedelta
from pathlib import Path

from flask import (Flask, render_template, request, jsonify,
                   send_file, abort, session, redirect, url_for)
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from functools import wraps
import hashlib, secrets, re

load_dotenv()

app = Flask(__name__, template_folder="ui/templates", static_folder="ui/static")
app.secret_key = os.environ.get("SECRET_KEY") or secrets.token_hex(32)
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"]  = "Strict"
app.config["SESSION_COOKIE_SECURE"]   = os.environ.get("HTTPS", "false").lower() == "true"
app.config["MAX_CONTENT_LENGTH"]      = 50 * 1024 * 1024   # 50MB max upload
from datetime import timedelta
app.permanent_session_lifetime = timedelta(hours=12)

BASE_DIR = Path(__file__).parent
REPORTS_DIR = BASE_DIR / "reports"
KNOWLEDGE_DIR = BASE_DIR / "knowledge"
TEMPLATES_DIR = BASE_DIR / "custom_templates"
LOGOS_DIR = BASE_DIR / "ui" / "static" / "logos"
DB_PATH = BASE_DIR / "database" / "penreport.db"

for d in [REPORTS_DIR, KNOWLEDGE_DIR, TEMPLATES_DIR, LOGOS_DIR, DB_PATH.parent]:
    d.mkdir(parents=True, exist_ok=True)


# ── Brute force protection ────────────────────────────────────────────────────
_failed_attempts = {}   # key (ip or username) -> [timestamp, ...]
_LOCKOUT_MAX     = 10   # max attempts per key
_LOCKOUT_WINDOW  = 300  # seconds (5 min)

def _check_rate_limit(key):
    """Returns True if key (IP or username) is rate-limited."""
    now = datetime.now().timestamp()
    attempts = [t for t in _failed_attempts.get(key, []) if now - t < _LOCKOUT_WINDOW]
    _failed_attempts[key] = attempts
    return len(attempts) >= _LOCKOUT_MAX

def _record_failed(key):
    now = datetime.now().timestamp()
    _failed_attempts.setdefault(key, []).append(now)

def _clear_failed(key):
    _failed_attempts.pop(key, None)


# ── Password helpers (must be before init_db) ────────────────────────────────

def _hash_password(password: str) -> str:
    salt = secrets.token_hex(16)
    h = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), 260000)
    return f"{salt}:{h.hex()}"

def _verify_password(password: str, stored: str) -> bool:
    try:
        salt, h = stored.split(":", 1)
        check = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), 260000)
        return secrets.compare_digest(h, check.hex())
    except Exception:
        return False


# ── Database ──────────────────────────────────────────────────────────────────

def get_db():
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    with get_db() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS reports (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                report_type TEXT NOT NULL,
                client TEXT,
                tester TEXT,
                scope TEXT,
                status TEXT DEFAULT 'draft',
                findings_raw TEXT,
                generated_content TEXT,
                template_used TEXT,
                target_system TEXT DEFAULT '',
                tested_from TEXT DEFAULT '',
                assessment_team TEXT DEFAULT '',
                created_by TEXT DEFAULT 'admin',
                visibility TEXT DEFAULT 'team',
                created_at TEXT DEFAULT (datetime('now')),
                updated_at TEXT DEFAULT (datetime('now'))
            )
        """)
        # Users table
        conn.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                role TEXT DEFAULT 'analyst',
                full_name TEXT DEFAULT '',
                created_at TEXT DEFAULT (datetime('now')),
                last_login TEXT
            )
        """)
        # report_shares table for per-user access control
        conn.execute("""
            CREATE TABLE IF NOT EXISTS report_shares (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                report_id INTEGER NOT NULL,
                shared_with TEXT NOT NULL,
                shared_by TEXT NOT NULL,
                created_at TEXT DEFAULT (datetime('now')),
                UNIQUE(report_id, shared_with)
            )
        """)
        # Migrations — safe to run on existing DB
        for migration in [
            "ALTER TABLE reports ADD COLUMN created_by TEXT DEFAULT 'admin'",
            "ALTER TABLE reports ADD COLUMN visibility TEXT DEFAULT 'team'",
            "ALTER TABLE users ADD COLUMN is_disabled INTEGER DEFAULT 0",
            # Phishing Campaign fields — added v1.4
            "ALTER TABLE reports ADD COLUMN reviewer TEXT DEFAULT ''",
            "ALTER TABLE reports ADD COLUMN approver TEXT DEFAULT ''",
            "ALTER TABLE reports ADD COLUMN campaign_period TEXT DEFAULT ''",
            "ALTER TABLE reports ADD COLUMN platform TEXT DEFAULT ''",
            "ALTER TABLE reports ADD COLUMN campaigns_json TEXT DEFAULT ''",
            "ALTER TABLE reports ADD COLUMN distribution_list TEXT DEFAULT ''",
            "ALTER TABLE reports ADD COLUMN phished_employees TEXT DEFAULT ''",
            "ALTER TABLE reports ADD COLUMN engagement_ref TEXT DEFAULT ''",
            "ALTER TABLE reports ADD COLUMN classification TEXT DEFAULT ''",
            "ALTER TABLE reports ADD COLUMN report_version TEXT DEFAULT ''",
        ]:
            try:
                conn.execute(migration)
            except Exception:
                pass
        conn.commit()

    # Create default admin if no users exist, or sync password if ADMIN_PASSWORD env var is set
    with get_db() as conn:
        count = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]
        env_pass = os.environ.get("ADMIN_PASSWORD")  # None if not in .env
        if count == 0:
            default_pass = env_pass or "vajra-admin-2026"
            _pw_hash = _hash_password(default_pass)
            conn.execute(
                "INSERT INTO users (username,password_hash,role,full_name) VALUES (?,?,?,?)",
                ("admin", _pw_hash, "admin", "Administrator")
            )
            conn.commit()
            print(f"\n  [TEAM] Default admin created: username=admin")
            print(f"  [TEAM] Password set from ADMIN_PASSWORD env var (see .env)")
            print(f"  [TEAM] Change the default password after first login!\n")
        elif env_pass and env_pass != "vajra-admin-2026":
            # Admin exists but .env has a custom ADMIN_PASSWORD — sync it every boot
            _pw_hash = _hash_password(env_pass)
            conn.execute("UPDATE users SET password_hash=? WHERE username='admin'", (_pw_hash,))
            conn.commit()
            print(f"\n  [TEAM] Admin password synced from ADMIN_PASSWORD env var.\n")

init_db()


# ── Report Access Control ─────────────────────────────────────────────────────

def can_access_report(report, username, role):
    """
    Check if a user can access a report.
    Access granted if ANY of:
    - user is admin
    - user created the report
    - visibility is 'public' (everyone)
    - visibility is 'team' AND role is admin or analyst (not viewer)
    - user is in report_shares for this report (explicit share)
    """
    if role == "admin":
        return True
    if report.get("created_by") == username:
        return True
    vis = report.get("visibility", "team")
    if vis == "public":
        return True
    if vis == "team" and role in ("admin", "analyst"):
        return True
    # Check shares table
    try:
        with get_db() as conn:
            share = conn.execute(
                "SELECT id FROM report_shares WHERE report_id=? AND shared_with=?",
                (report.get("id"), username)
            ).fetchone()
            return share is not None
    except Exception:
        return False

def can_modify_report(report, username, role):
    """Can delete/change visibility — owner or admin only."""
    return role == "admin" or report.get("created_by") == username


# ── Knowledge Base ────────────────────────────────────────────────────────────

# ── Knowledge Base Selection Map ─────────────────────────────────────────────
# Maps assessment signals → relevant KB files.
# kb-core-reference.md is ALWAYS loaded (severity/CVSS/CWE/remediation).
# Domain-specific files are loaded based on report type + keywords in scope/findings.

_KB_ALWAYS = ["kb-core-reference.md", "severity-definitions.md", "cwe-reference.md", "remediation-boilerplate.md"]

# Signals use word-boundary matching (regex \b) to avoid false positives.
# e.g. "ad" won't match "credentials", "domain" won't match "random".
# Multi-word phrases are matched as-is (spaces already act as boundaries).
# Overly generic single-word tokens removed from each category.
_KB_SIGNAL_MAP = {
    # Web application — require explicit web/app terms
    "kb-web-application.md": [
        "web application", "web app", "webapp", "web portal", "web server",
        "owasp", "xss", "sql injection", "csrf", "ssrf", "xxe",
        "http response", "cookie", "session token", "web login", "web portal",
        "dashboard", "api gateway", "graphql", "browser exploit", "frontend",
        "web vulnerability", "web security", "web pentest",
    ],
    # API security — require explicit API terms
    "kb-api-security.md": [
        "rest api", "api endpoint", "graphql api", "soap api",
        "swagger", "openapi", "oauth", "jwt token", "api security",
        "microservice", "webhook", "grpc", "api testing", "api key",
        "postman", "api authentication", "broken api",
    ],
    # Active Directory — require specific AD terms, not generic "domain" or "windows"
    "kb-active-directory.md": [
        "active directory", "domain controller", "domain admin",
        "kerberos", "kerberoast", "ntlm", "ldap", "smb relay",
        "pass the hash", "pass-the-hash", "bloodhound", "mimikatz",
        "group policy", "gpo", "lsass", "ntds", "azure ad", "entra id",
        "lateral movement", "privilege escalation", "ad environment",
        "internal network", "windows domain",
    ],
    # Cloud security
    "kb-cloud-security.md": [
        "cloud", "aws", "azure", "gcp", "google cloud", "s3 bucket",
        "ec2", "iam role", "kubernetes", "k8s", "docker container",
        "serverless", "lambda function", "terraform", "cloudformation",
        "blob storage", "cloud security", "cloud misconfiguration",
        "eks", "aks", "gke", "cloud trail",
    ],
    # Network infrastructure — require specific network terms
    "kb-network-infrastructure.md": [
        "network infrastructure", "firewall", "router", "vpn", "dmz",
        "port scan", "nmap", "nessus", "vulnerability scan",
        "network perimeter", "external perimeter", "dns server",
        "snmp", "telnet", "network segmentation", "packet capture",
        "network pentest", "infrastructure assessment",
    ],
    # Mobile security
    "kb-mobile-security.md": [
        "mobile app", "mobile application", "android", "ios app",
        "apk", "ipa", "app store", "google play", "react native",
        "flutter", "thick client", "desktop app", "electron app",
        "mobile security", "mobile pentest",
    ],
    # Wireless / IoT — require specific wireless/IoT terms
    "kb-wireless-iot-misc.md": [
        "wireless", "wifi", "wi-fi", "wpa2", "wpa3", "bluetooth",
        "iot device", "embedded device", "firmware", "uart", "jtag",
        "zigbee", "rfid", "nfc tag", "rogue ap", "evil twin",
        "deauth attack", "wireless security", "iot security",
    ],
    # Red team — require explicit red team terms
    "kb-red-team.md": [
        "red team", "redteam", "apt simulation", "threat actor",
        "cobalt strike", "havoc", "sliver", "metasploit", "c2 beacon",
        "command and control", "persistence mechanism", "data exfiltration",
        "evasion technique", "av bypass", "edr bypass", "opsec",
        "mitre att&ck", "kill chain", "assumed breach", "purple team",
        "adversary emulation", "post-exploitation",
    ],
    # Phishing / social engineering
    "kb-phishing-social-engineering.md": [
        "phishing", "spear phishing", "vishing", "smishing", "bec",
        "business email compromise", "social engineering", "pretext",
        "credential harvest", "credentials harvested", "harvested credentials",
        "aitm", "evilginx", "gophish",
        "phishing awareness", "phishing campaign", "email security",
        "dmarc", "spf", "dkim", "phishing simulation",
        "user awareness", "phishing template", "click rate",
        "clicked the link", "submitted credentials", "mfa bypass",
    ],
    # OWASP Top 10
    "owasp-top10.md": [
        "owasp top 10", "owasp top ten", "broken access control",
        "cryptographic failure", "injection vulnerability",
        "insecure design", "security misconfiguration",
        "vulnerable component", "identification failure",
        "software integrity", "logging failure", "server-side request",
    ],
}

def select_kb_files(report_type: str, scope: str = "", findings_text: str = "",
                    target_system: str = "") -> list:
    """
    Intelligently select relevant KB files based on assessment context.
    Uses word-boundary regex matching to avoid false positives from short tokens.
    Always includes core reference files. Domain files selected only on clear signals.
    """
    import re as _re2

    # Combine all available context text for signal detection
    context = " ".join([
        (report_type or "").lower(),
        (scope or "").lower(),
        (findings_text or "")[:2000].lower(),
        (target_system or "").lower(),
    ])

    selected = list(_KB_ALWAYS)

    def _matches(signal: str, ctx: str) -> bool:
        """Word-boundary aware match. Multi-word phrases use direct search."""
        if " " in signal:
            return signal in ctx
        # Single word: require word boundary to avoid "ad" in "loaded", "web" in "webhook"
        return bool(_re2.search(r'' + _re2.escape(signal) + r'', ctx))

    # Match domain-specific files based on signals in context
    for kb_file, signals in _KB_SIGNAL_MAP.items():
        if any(_matches(s, context) for s in signals):
            if kb_file not in selected:
                selected.append(kb_file)

    # Report-type overrides — always include type-specific KB
    if report_type == "redteam":
        for f in ["kb-red-team.md", "kb-active-directory.md"]:
            if f not in selected:
                selected.append(f)
    elif report_type == "phishing":
        if "kb-phishing-social-engineering.md" not in selected:
            selected.append("kb-phishing-social-engineering.md")
    elif report_type == "pentest":
        # Default pentest: add web + network if nothing specific matched
        if not any(f in selected for f in [
            "kb-web-application.md", "kb-api-security.md",
            "kb-network-infrastructure.md", "kb-active-directory.md",
            "kb-mobile-security.md", "kb-cloud-security.md",
        ]):
            selected.append("kb-web-application.md")
            selected.append("kb-network-infrastructure.md")

    return selected


def load_knowledge_base(selected_only=None, report_type="pentest",
                        scope="", findings_text="", target_system=""):
    """
    Intelligently load relevant KB files based on assessment context.
    Always loads core reference files. Domain files selected by signal matching.
    selected_only: if set, only load that specific filename (legacy support).
    """
    # Legacy single-file mode
    if selected_only:
        try:
            fp = KNOWLEDGE_DIR / selected_only
            content = fp.read_text(encoding="utf-8")
            print(f"[KB] Loaded (specific): {selected_only} ({len(content)} chars)")
            return f"=== KNOWLEDGE: {selected_only} ===\n{content}"
        except Exception as e:
            print(f"[KB] Failed to load {selected_only}: {e}")
            return f"Knowledge file {selected_only} not found."

    # Smart selection
    to_load = select_kb_files(report_type, scope, findings_text, target_system)
    print(f"[KB] Smart selection for [{report_type}]: {to_load}")

    parts = []
    skipped = []
    for fname in to_load:
        fp = KNOWLEDGE_DIR / fname
        if not fp.exists():
            # Also try glob for user-uploaded files not in map
            skipped.append(fname)
            continue
        try:
            content = fp.read_text(encoding="utf-8")
            parts.append(f"=== KNOWLEDGE: {fname} ===\n{content}")
            print(f"[KB] Loaded: {fname} ({len(content)} chars)")
        except Exception as e:
            print(f"[KB] Failed to load {fname}: {e}")

    # Also load any user-uploaded .md/.txt files NOT in the map
    # (custom KB files the user has added)
    map_files = set(_KB_ALWAYS) | set(_KB_SIGNAL_MAP.keys())
    for fp in sorted(KNOWLEDGE_DIR.glob("*.md")) + sorted(KNOWLEDGE_DIR.glob("*.txt")):
        fname = fp.name
        if fname not in map_files and fname not in to_load:
            try:
                content = fp.read_text(encoding="utf-8")
                parts.append(f"=== KNOWLEDGE: {fname} ===\n{content}")
                print(f"[KB] Loaded (custom): {fname} ({len(content)} chars)")
            except Exception as e:
                print(f"[KB] Failed to load {fname}: {e}")

    if not parts:
        return "No knowledge base files found."

    total = sum(len(p) for p in parts)
    print(f"[KB] Total: {len(parts)} files, {total} chars "
          f"(skipped: {skipped if skipped else 'none'})")
    return "\n\n".join(parts)


# ── Template Parsing ──────────────────────────────────────────────────────────

def parse_docx_structure(filepath):
    try:
        from docx import Document
        doc = Document(str(filepath))
        structure = []
        for para in doc.paragraphs:
            if para.text.strip():
                structure.append({"style": para.style.name, "text": para.text.strip()})
        return structure
    except Exception as e:
        return [{"style": "Normal", "text": f"Could not parse: {e}"}]

def parse_md_structure(filepath):
    structure = []
    try:
        with open(str(filepath), "r", encoding="utf-8") as f:
            for line in f:
                line = line.rstrip()
                if line.startswith("#"):
                    level = len(line) - len(line.lstrip("#"))
                    structure.append({"style": f"Heading {level}", "text": line.lstrip("# ").strip()})
                elif line.strip():
                    structure.append({"style": "Normal", "text": line.strip()[:80]})
    except Exception as e:
        structure.append({"style": "Normal", "text": f"Error: {e}"})
    return structure

def inject_placeholders(content, meta):
    """
    Replace {{placeholder}} tokens in template content with real engagement data.
    Validates and logs missing values.
    """
    from datetime import datetime as _dt
    import re as _re

    date_str = meta.get("date") or _dt.now().strftime("%Y-%m-%d")
    eng_ref  = f"ENG-{date_str[:4]}-001"

    mapping = {
        # Long-form keys
        "client_name":        meta.get("client")          or "Client",
        "prepared_by":        meta.get("tester")          or "Security Assessor",
        "assessment_team":    meta.get("assessment_team") or meta.get("tester") or "Security Team",
        "report_date":        date_str,
        "engagement_ref":     eng_ref,
        "target_system":      meta.get("target_system")   or meta.get("scope") or "As defined in RoE",
        "test_scope":         meta.get("scope")           or "As defined in Rules of Engagement",
        "engagement_scope":   meta.get("scope")           or "As defined in Rules of Engagement",
        "tested_from":        meta.get("tested_from")     or "External Network",
        "classification":     "CONFIDENTIAL — NOT FOR DISTRIBUTION",
        # Short-form aliases — match what the VAJRA templates actually use
        "client":             meta.get("client")          or "Client",
        "tester":             meta.get("tester")          or "Security Assessor",
        "author":             meta.get("author")          or meta.get("tester") or "Security Assessor",
        "date":               date_str,
        "scope":              meta.get("scope")           or "As defined in Rules of Engagement",
        "assessment_type":    meta.get("assess_type")     or "External Black-Box Penetration Test",
        "report_title":       f"{(meta.get('report_type') or 'Penetration Test').title()} Report",
        "lead_assessor":      meta.get("tester")          or "Security Assessor",
        "client_org":         meta.get("client")          or "Client",
        "version":            meta.get("report_version")  or "1.0 — Final",
        "report_version":     "1.0 — Final",
        "report_title":       "Security Assessment Report",
        "testing_methodology":"OWASP WSTG v4.2, PTES, CVSS v3.1, NIST SP 800-115",
        "executive_summary":  "[AI will generate executive summary based on findings]",
        "findings_summary_table": "[AI will generate findings summary table]",
        "detailed_findings":  "[AI will generate detailed findings]",
        "remediation_roadmap":"[AI will generate remediation roadmap]",
        "tools_used":         "Burp Suite Pro, Nmap, OWASP ZAP, Nikto, sqlmap",
        "references":         "OWASP Top 10:2021, CVSS v3.1, NIST SP 800-115",
    }

    # Find all placeholders in template and log missing ones
    # Ignore meta-words that appear in template documentation, not as real placeholders
    _meta_words = {"placeholder", "placeholders", "field", "fields", "value", "example"}
    found = _re.findall(r'\{\{(\w+)\}\}', content)
    for ph in set(found):
        if ph not in mapping and ph not in _meta_words:
            print(f"[TEMPLATE] INFO: No data for {{{{{{ph}}}}}} — will be passed to AI as-is")

    # Replace all known placeholders
    for key, val in mapping.items():
        content = content.replace(f"{{{{{key}}}}}", val)

    injected = set(found) & set(mapping.keys())
    print(f"[TEMPLATE] Injected {len(injected)} placeholder(s): {sorted(injected)}")
    return content


def build_template_context(template_name, meta=None):
    """
    Issue 1 fix: Load the FULL selected template content (not just headings).
    Issue 6 fix: Replace all {{placeholder}} tokens with real engagement data.
    The result is passed to the AI as the authoritative structure to follow.
    """
    if not template_name:
        print("[TEMPLATE] No template selected — AI will use default structure")
        return "No specific template selected — use standard professional pentest report structure."

    filepath = TEMPLATES_DIR / template_name
    if not filepath.exists():
        print(f"[TEMPLATE] WARNING: File not found: {filepath}")
        return f"Template '{template_name}' not found — use standard structure."

    ext = filepath.suffix.lower()
    print(f"[TEMPLATE] Loading: {template_name} ({ext})")

    if ext in [".md", ".txt"]:
        try:
            with open(str(filepath), "r", encoding="utf-8") as f:
                raw = f.read()
            # Inject placeholder values so AI sees real data
            content = inject_placeholders(raw, meta or {})
            print(f"[TEMPLATE] Loaded {len(content)} chars — passing full content to AI")
            return (
                f"SELECTED TEMPLATE: {template_name}\n"
                f"Follow this template structure EXACTLY. Preserve all sections and headings.\n"
                f"Replace any remaining placeholders with appropriate generated content.\n"
                f"\n--- TEMPLATE START ---\n{content}\n--- TEMPLATE END ---"
            )
        except Exception as e:
            print(f"[TEMPLATE] Error: {e}")
            return f"Could not read template '{template_name}': {e}"

    elif ext == ".docx":
        structure = parse_docx_structure(filepath)
        lines = []
        for item in structure:
            style = item.get("style", "")
            text  = item.get("text", "")
            if not text or text.startswith("${"):
                continue
            if "Title" in style or "Heading 1" in style:
                lines.append(f"\n# {text}")
            elif "Heading 2" in style:
                lines.append(f"## {text}")
            elif "Heading 3" in style:
                lines.append(f"### {text}")
            else:
                lines.append(f"  {text[:120]}")
        out = "\n".join(lines) or "Use standard structure."
        print(f"[TEMPLATE] Loaded DOCX: {len(lines)} sections")
        return f"SELECTED TEMPLATE (DOCX): {template_name}\nFollow this structure:\n{out}"

    else:
        return f"Unsupported template format: {ext}"


# ── Tool Import Parsers ───────────────────────────────────────────────────────

def parse_nessus(xml_content):
    findings = []
    try:
        root = ET.fromstring(xml_content)
        sev_map = {"0": "Info", "1": "Low", "2": "Medium", "3": "High", "4": "Critical"}
        for host in root.iter("ReportHost"):
            hostname = host.get("name", "Unknown")
            for item in host.iter("ReportItem"):
                findings.append({
                    "title": item.get("pluginName", "Unknown"),
                    "severity": sev_map.get(item.get("severity", "0"), "Info"),
                    "host": hostname,
                    "port": item.get("port", ""),
                    "description": item.findtext("description", ""),
                    "solution": item.findtext("solution", ""),
                    "cve": item.findtext("cve", ""),
                    "cvss": item.findtext("cvss3_base_score", item.findtext("cvss_base_score", ""))
                })
    except Exception as e:
        findings.append({"title": f"Parse error: {e}", "severity": "Info"})
    return findings

def parse_burp(xml_content):
    findings = []
    try:
        root = ET.fromstring(xml_content)
        for issue in root.iter("issue"):
            findings.append({
                "title": (issue.findtext("name") or "").strip(),
                "severity": (issue.findtext("severity") or "Info").strip(),
                "host": (issue.findtext("host") or "").strip(),
                "path": (issue.findtext("path") or "").strip(),
                "description": (issue.findtext("issueDetail") or issue.findtext("issueBackground") or "").strip(),
                "remediation": (issue.findtext("remediationDetail") or issue.findtext("remediationBackground") or "").strip(),
            })
    except Exception as e:
        findings.append({"title": f"Parse error: {e}", "severity": "Info"})
    return findings

def parse_nmap(xml_content):
    findings = []
    try:
        root = ET.fromstring(xml_content)
        for host in root.iter("host"):
            addr = host.find("address")
            ip = addr.get("addr", "unknown") if addr is not None else "unknown"
            hn = host.find(".//hostname")
            hostname = hn.get("name", ip) if hn is not None else ip
            for port in host.iter("port"):
                state = port.find("state")
                if state is not None and state.get("state") == "open":
                    portid = port.get("portid", "")
                    proto = port.get("protocol", "tcp")
                    svc = port.find("service")
                    svc_name = f"{svc.get('name','')} {svc.get('product','')}".strip() if svc is not None else ""
                    findings.append({
                        "title": f"Open Port {portid}/{proto} — {svc_name}".strip(" —"),
                        "severity": "Info",
                        "host": hostname,
                        "port": portid,
                        "description": f"Host {hostname} has port {portid}/{proto} open. {svc_name}"
                    })
    except Exception as e:
        findings.append({"title": f"Parse error: {e}", "severity": "Info"})
    return findings

def findings_to_text(findings):
    lines = [f"IMPORTED FINDINGS ({len(findings)} total)\n"]
    for i, f in enumerate(findings, 1):
        lines.append(f"Finding #{i}:")
        for k, v in f.items():
            if v:
                lines.append(f"  {k.title()}: {str(v)[:500]}")
        lines.append("")
    return "\n".join(lines)


# ── AI Report Generation ──────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are a CREST-certified senior penetration tester writing professional security assessment reports.
Standards: CREST / PTES / OWASP WSTG v4.2 / NIST SP 800-115 / CVSS v3.1.
STRICT RULES — never break these:
- Output ONLY clean Markdown. No preamble, no commentary, no "Here is your report:" intro.
- NEVER emit emoji for severity. Use plain text: Critical / High / Medium / Low / Info.
- NEVER emit markdown anchor links [text](#anchor) — they render as raw text in Word.
- NEVER output --- horizontal rules as section separators.
- NEVER output a Table of Contents, cover page, Document Control, Revision History, or second Findings Overview — the tool generates these automatically.
- All finding IDs: F-001, F-002, F-003 (3-digit zero-padded) everywhere consistently.
- Remediation deadlines: calculate actual calendar dates (Critical+2d, High+7d, Medium+30d, Low+90d from the report date).
- Every finding must have Retest Status: Pending Retest.
COMPLETENESS — CRITICAL RULE:
- You MUST generate a COMPLETE ### F-NNN detailed block for EVERY single finding listed in the summary table.
- If the summary table shows 5 findings (F-001 through F-005), you MUST generate all 5 detailed blocks.
- Do NOT stop after 2 or 3 findings. Do NOT truncate. Complete ALL findings no matter how many.
- If you listed a finding in the Findings Summary table, it MUST have a full detailed block in Section 3.
- Each finding block MUST include: Description, Steps to Reproduce, Evidence, Business Impact, Remediation.
SECTION NAMES — use EXACTLY these names so the DOCX builder can parse them:
- Use "#### Description" (not "#### Overview" or "#### Details")
- Use "#### Steps to Reproduce" (not "#### PoC" or "#### Reproduction")
- Use "#### Evidence" (not "#### Proof of Concept" alone)
- Use "#### Business Impact" (not "#### Impact Analysis" or "#### Impact")
- Use "#### Remediation" (not "#### Remediation Guidance" or "#### Fix")
KNOWLEDGE BASE USAGE:
- The KNOWLEDGE BASE section contains CWE references, OWASP Top 10, severity definitions, and remediation boilerplate.
- You MUST use this knowledge to: enrich finding descriptions, assign accurate CWE IDs, map to OWASP categories, provide standard remediation guidance, and calibrate severity ratings.
TEMPLATE USAGE:
- If a SELECTED TEMPLATE is provided, follow its section order and structure EXACTLY.
- Replace any remaining {{placeholder}} tokens with appropriate generated content."""

REPORT_PROMPTS = {
    "pentest": """Generate a professional Penetration Test Report in Markdown (CREST/PTES/OWASP WSTG v4.2).

ENGAGEMENT:
- Client: {client}
- Lead Assessor: {tester} | Assessment Team: {assessment_team}
- Date: {date} | Target: {target_system} | Scope: {scope} | Tested From: {tested_from}

FINDINGS PROVIDED:
{findings}

KNOWLEDGE BASE:
{knowledge}

TEMPLATE CONTEXT:
{template_structure}

OUTPUT THESE SECTIONS IN ORDER — no extras, no skipped sections:

# Executive Summary
[2–3 paragraphs. Board-level, non-technical. State: what was tested, findings by severity count, top risk in one sentence, business impact, single most important action. Name {client} throughout.]

## Overall Risk Assessment
| Metric | Value |
|---|---|
| Critical Findings | [N] |
| High Findings | [N] |
| Medium Findings | [N] |
| Low Findings | [N] |
| Informational | [N] |
| Overall Risk | [Critical/High/Medium/Low] |
| Engagement Reference | ENG-{date}-001 |
| Report Version | 1.0 — Final |

# 1. Engagement Overview

## 1.1 Scope and Rules of Engagement
| Field | Details |
|---|---|
| Client | {client} |
| Target System | {target_system} |
| In-Scope | {scope} |
| Out-of-Scope | All systems not listed above |
| Assessment Type | External Black-Box Penetration Test |
| Tested From | {tested_from} |
| Testing Period | {date} |
| Lead Assessor | {tester} |
| Assessment Team | {assessment_team} |
| Classification | CONFIDENTIAL — NOT FOR DISTRIBUTION |

**In-Scope:** [bullet list from scope]
**Out-of-Scope:** Third-party integrations not on client infrastructure; DoS testing; social engineering

**Constraints:** Testing was time-boxed; findings are point-in-time; no production data was exfiltrated.

## 1.2 Methodology
Testing followed OWASP WSTG v4.2, PTES, CVSS v3.1, and NIST SP 800-115 across six phases: Reconnaissance → Scanning → Vulnerability Identification → Exploitation → Post-Exploitation → Reporting.

## 1.3 Risk Rating
| Severity | CVSS v3.1 | SLA | Definition |
|---|---|---|---|
| Critical | 9.0–10.0 | 24–48 hours | Trivial exploit, no prerequisites. Full compromise. |
| High | 7.0–8.9 | 7 days | Data breach or privilege escalation likely. |
| Medium | 4.0–6.9 | 30 days | Requires conditions or chaining. |
| Low | 0.1–3.9 | 90 days | Minimal standalone impact. |
| Info | 0.0 | Best effort | Hardening opportunity, no exploit. |

## 1.4 Limitations
- Time-boxed assessment — additional vulnerabilities may exist
- Scope-limited — out-of-scope systems were not tested
- Point-in-time — posture may have changed since testing
- Non-destructive — no production data was exfiltrated or modified

# 2. Findings Summary

## 2.1 Vulnerability Summary
| ID | Title | Severity | CVSS | Affected Asset | Retest Status |
|---|---|---|---|---|---|
[one row per finding, IDs F-001 F-002 etc.]

## 2.2 Severity Distribution
| Severity | Count | % of Total |
|---|---|---|
[one row per severity]

# 3. Detailed Findings

[For EVERY finding, use EXACTLY this format:]

### F-001 — [Title]

| Field | Detail |
|---|---|
| Severity | [Critical/High/Medium/Low/Info] |
| CVSS v3.1 Score | [N.N] ([rating]) |
| CVSS Vector | `AV:X/AC:X/PR:X/UI:X/S:X/C:X/I:X/A:X` |
| CWE | CWE-[ID]: [Full name] |
| CVE | [CVE-YYYY-NNNNN or: No CVE — Novel Finding] |
| OWASP Top 10 | A0X:2021 — [Category] |
| Affected Host | [URL or hostname] |
| Affected Parameter | [parameter or N/A] |
| Technology | [e.g. PHP 7.4 / MySQL 5.7] |
| Root Cause | [one sentence] |
| Likelihood | High / Medium / Low |
| Business Impact | Critical / High / Medium / Low |
| Remediation Deadline | [calculated calendar date from {date}] |
| Retest Status | Pending Retest |

#### Description
[2 paragraphs: what the vuln is, where it is, why it exists, how discovered]

#### Steps to Reproduce
1. [exact step]
2. [exact step]
3. [confirm exploitation indicator]

#### Evidence
[Brief context, then:]
```
[FULL HTTP REQUEST as sent]
```
[Brief observation, then:]
```
[FULL HTTP RESPONSE / OUTPUT]
```
> 📸 Screenshot/PoC: Insert screenshot evidence here

#### Business Impact
- **[Category]:** [specific impact referencing {client}'s data/systems]
- **Regulatory Exposure:** [GDPR/PCI DSS/FINRA clause if applicable]

#### Remediation
1. **Immediate:** [specific primary fix with code example if applicable]
2. **Defense-in-depth:** [compensating control]
3. **Code review scope:** [what else to audit]
4. **Verification:** [how to confirm fix worked]
Reference: [OWASP cheatsheet or CWE URL]

[Repeat ### F-001 block for EVERY finding. Do NOT skip any.]

# 4. Remediation Roadmap

## 4.1 Prioritised Schedule
| # | ID | Title | Severity | Deadline | Owner | Status |
|---|---|---|---|---|---|---|
[one row per finding with calculated calendar deadlines from {date}]

## 4.2 Remediation Phases
**Phase 1 — Emergency (0–48 hours):** [Critical findings + first action each]
**Phase 2 — Urgent (7 days):** [High findings + first action each]
**Phase 3 — Planned (30 days):** [Medium findings + first action each]
**Phase 4 — Backlog (90 days):** [Low findings + first action each]
**Phase 5 — Ongoing:** Schedule annual pentest; integrate DAST into CI/CD; establish vuln management programme.

## 4.3 Verification Process
Upon remediation notification, the assessment team will retest within 10 business days and issue a Remediation Verification Letter. Report updated as Version 1.1.

# 5. Appendix

## Appendix A — Tools Used
| Tool | Purpose |
|---|---|
| Burp Suite Pro | HTTP interception, manual testing, active scanning |
| OWASP ZAP | Automated scanning |
| Nmap | Port scanning, service detection |
| Nikto | Web server misconfiguration |
| sqlmap | SQL injection confirmation |

## Appendix B — Regulatory Mapping
| Regulation | Requirement | Findings |
|---|---|---|
| PCI DSS v4.0 | Req. 11.3: Penetration testing | All |
| GDPR Article 32 | Technical security measures | [findings exposing PII] |
| ISO 27001:2022 | A.8.8: Technical vulnerability management | All |

## Appendix C — Legal Disclaimer
This assessment was conducted under written authorisation from {client}. This report is CONFIDENTIAL — NOT FOR DISTRIBUTION. Findings represent a point-in-time assessment. {client} is solely responsible for remediation decisions.

## Appendix D — References
- OWASP Top 10:2021: https://owasp.org/Top10/
- CVSS v3.1 Calculator: https://www.first.org/cvss/calculator/3-1
- OWASP Testing Guide v4.2: https://owasp.org/www-project-web-security-testing-guide/
- CWE Top 25: https://cwe.mitre.org/top25/
- NIST NVD: https://nvd.nist.gov/""",

    "redteam": """Generate a professional Red Team Engagement Report (CREST STAR / CBEST / TIBER-EU standards).

ENGAGEMENT:
- Client: {client}
- Lead Assessor: {tester} | Team: {assessment_team}
- Date: {date} | Scope: {scope} | Tested From: {tested_from}

ENGAGEMENT NOTES AND FINDINGS:
{findings}

KNOWLEDGE BASE:
{knowledge}

TEMPLATE CONTEXT:
{template_structure}

OUTPUT THESE SECTIONS:

# Executive Summary
[2–3 paragraphs. Board-level. State: objective, threat scenario simulated, whether crown jewels reached, key risk indicators, top defensive recommendation. Name {client}.]

## Engagement Metrics
| Metric | Value |
|---|---|
| Attack Scenarios | [N] |
| Objectives Achieved | [N of N] |
| Crown Jewels Reached | Yes / No |
| Blue Team Detection Rate | [X%] |
| Mean Time to Detect | [X hours or Not Detected] |
| Report Version | 1.0 — Final |

# 1. Engagement Overview

## 1.1 Objectives and Crown Jewels
**Primary Objective:** [What the red team was tasked to achieve]
**Crown Jewels:** [High-value target assets]
**Threat Scenario:** [APT profile or adversary simulated]

## 1.2 Scope
| Field | Details |
|---|---|
| Client | {client} |
| In-Scope | {scope} |
| Type | Red Team — Simulated External Adversary |
| Tested From | {tested_from} |
| Date | {date} |
| Lead Assessor | {tester} |
| Classification | CONFIDENTIAL — RESTRICTED DISTRIBUTION |

## 1.3 Kill Chain Methodology
| Phase | MITRE Tactic | Description |
|---|---|---|
| Reconnaissance | TA0043 | OSINT, infrastructure mapping |
| Initial Access | TA0001 | External exploitation, credential attacks |
| Execution | TA0002 | Payload delivery |
| Persistence | TA0003 | Long-term access mechanisms |
| Privilege Escalation | TA0004 | Elevated system/domain access |
| Lateral Movement | TA0008 | Pivoting through network |
| Collection/Exfiltration | TA0009/TA0010 | Achieving objectives |

# 2. Attack Narrative
[Chronological prose. Full kill chain story with timestamps, specific techniques, and pivots. Most important section for client understanding.]

## 2.1 Phase-by-Phase Breakdown
### Reconnaissance
[OSINT gathered, intelligence used for planning]
### Initial Access
[Successful vector, what was exploited, timeline]
### Persistence and Privilege Escalation
[Mechanisms deployed, credentials harvested]
### Lateral Movement and Objective Achievement
[Systems pivoted through, crown jewels reached or not, why]

# 3. Observations and Findings

[For each observation:]
### O-001 — [Title]
| Field | Details |
|---|---|
| MITRE Tactic | [TA00XX — Name] |
| MITRE Technique | [TXXXX — Name] |
| Severity | [Critical/High/Medium/Low] |
| Affected System | [system] |
| Detected by Blue Team | Yes / No / Partial |
| Retest Status | Pending Retest |

#### Description
[Technical detail]

#### Evidence
```
[Evidence in monospace]
```
> 📸 Screenshot/PoC: Insert screenshot evidence here

#### Impact
[Business and technical impact for {client}]

#### Recommendations
1. [Specific fix]
2. [Defense-in-depth]

# 4. MITRE ATT&CK Mapping
| Tactic | Technique ID | Technique | Phase |
|---|---|---|---|
[one row per TTP used]

# 5. Detection and Response Assessment

## 5.1 Blue Team Performance
| Phase | Activity | Detected | Time to Detect | Response |
|---|---|---|---|---|
[one row per major activity]

## 5.2 Detection Gaps
[What SOC missed and why — logging gaps, tuning issues, blind spots]

## 5.3 Improvement Recommendations
[Specific SIEM rules, log sources, monitoring improvements]

# 6. Remediation Roadmap
| # | Observation | Severity | Deadline | Owner | Action |
|---|---|---|---|---|---|
[calculated dates from {date}]

# 7. Appendix
## Appendix A — Tools Used
| Tool | Purpose | Phase |
|---|---|---|

## Appendix B — Legal Disclaimer
Conducted under written authorisation from {client}. CONFIDENTIAL — RESTRICTED DISTRIBUTION.""",

    "phishing": """Generate a professional Phishing Campaign Assessment Report (NIST / SANS / CREST standards).

ENGAGEMENT:
- Client: {client}
- Lead Assessor: {tester} | Team: {assessment_team}
- Date: {date} | Target Population: {scope}

CAMPAIGN DATA:
{findings}

KNOWLEDGE BASE:
{knowledge}

TEMPLATE CONTEXT:
{template_structure}

OUTPUT THESE SECTIONS:

# Executive Summary
[2–3 paragraphs. Human risk and organisational exposure focus. Quantify susceptibility rate. Name {client}. State single most important awareness action.]

## Campaign Metrics
| Metric | Count | Rate |
|---|---|---|
| Emails Sent | [N] | 100% |
| Emails Delivered | [N] | [X%] |
| Emails Opened | [N] | [X%] |
| Links Clicked | [N] | [X%] |
| Credentials Submitted | [N] | [X%] |
| MFA Bypassed | [N] | [X%] |
| Reported by Users | [N] | [X%] |
| Mean Time to First Click | [X minutes] | — |

# 1. Engagement Overview

## 1.1 Campaign Configuration
| Field | Details |
|---|---|
| Client | {client} |
| Campaign Date | {date} |
| Assessment Team | {assessment_team} |
| Target Population | {scope} |
| Pretext / Lure | [lure type] |
| Sending Domain | [domain used] |
| Payload Type | Credential Harvest / Malware / Awareness Only |
| Classification | CONFIDENTIAL — NOT FOR DISTRIBUTION |

## 1.2 Authorisation
Testing conducted under written authorisation from {client}. No actual malicious payloads deployed. All simulated credential submissions immediately invalidated.

# 2. Results Analysis

## 2.1 Overall Results
[Summary with industry benchmark — average click rate ~17% (Proofpoint SBIR 2023)]

## 2.2 Department Breakdown
| Department | Sent | Clicked | Submitted | Click Rate | Risk Level |
|---|---|---|---|---|---|
[one row per dept if data available]

## 2.3 Email Security Controls
| Control | Configured | Effective | Bypassed |
|---|---|---|---|
| SPF | Yes/No | Yes/No | Yes/No |
| DKIM | Yes/No | Yes/No | Yes/No |
| DMARC | Yes/No | Yes/No | Yes/No |
| Email Gateway | Yes/No | Yes/No | Yes/No |

# 3. Observations and Findings

### P-001 — [Finding Title]
| Field | Details |
|---|---|
| Severity | [Critical/High/Medium/Low] |
| Affected Group | [Dept/role] |
| Click Rate | [X%] |
| Submission Rate | [X%] |
| Root Cause | [Training gap / No MFA / Permissive gateway] |
| Retest Status | Pending Retest |

#### Description
[Statistical and contextual detail]

#### Evidence
> 📸 Screenshot/PoC: Insert screenshot evidence here

#### Impact
[Business risk if exploited by real threat actor]

#### Recommendations
1. [Specific fix for this group]
2. [Technical control]

# 4. Recommendations

## 4.1 Immediate Actions (before next business day)
1. Reset credentials for all users who submitted during the campaign
2. Enforce MFA on all accounts that submitted credentials
3. [Additional immediate actions]

## 4.2 Technical Controls
| Control | Priority | Action |
|---|---|---|
| DMARC enforcement | [H/M/L] | [specific action] |
| MFA deployment | [H/M/L] | [specific action] |
| Email gateway tuning | [H/M/L] | [specific action] |

## 4.3 Security Awareness Programme
[Targeted training recommendations by risk group. Reference specific platforms — KnowBe4, Proofpoint SA, SANS OUCH.]

# 5. Remediation Roadmap
| # | Finding | Severity | Deadline | Owner | Action |
|---|---|---|---|---|---|
[calculated dates from {date}]

# 6. Appendix
## Appendix A — Legal Disclaimer
Conducted under authorisation from {client}. No malicious payloads deployed. CONFIDENTIAL.

## Appendix B — References
- Proofpoint State of the Phish: https://www.proofpoint.com/us/resources/threat-reports/state-of-phish
- NIST SP 800-177: https://csrc.nist.gov/publications/detail/sp/800/177/final
- Anti-Phishing Working Group: https://apwg.org/"""
}


# ── AI Provider Router ────────────────────────────────────────────────────────
# Zero extra installs required for most providers.
# httpx is already a core dependency — all OpenAI-compatible providers use it
# directly via raw HTTP calls (no openai SDK needed).
#
# Only two providers need an optional extra package:
#   gemini  → pip3 install google-generativeai   (Google's SDK has no REST alternative)
#   mistral → pip3 install mistralai             (streaming-first SDK, REST is awkward)
#   cohere  → pip3 install cohere                (proprietary request format)
#
# All others (OpenAI, Grok, Groq, Together, Ollama, DeepSeek, OpenRouter,
# Perplexity, LM Studio, vLLM, LiteLLM, Azure OpenAI, Anthropic) work out
# of the box with zero extra installs.
#
# To add a new provider:
#   1. If it uses /v1/chat/completions  → add its URL to _OPENAI_COMPAT_URLS
#   2. If it uses /v1/messages (Anthropic spec) → add to _ANTHROPIC_COMPAT_URLS
#   3. If it has a unique API format  → add an elif block in _call_provider()

# OpenAI /v1/chat/completions compatible — just a base URL, no extra package.
_OPENAI_COMPAT_URLS = {
    "openai":      "https://api.openai.com/v1",
    "chatgpt":     "https://api.openai.com/v1",
    "grok":        "https://api.x.ai/v1",
    "groq":        "https://api.groq.com/openai/v1",
    "together":    "https://api.together.xyz/v1",
    "ollama":      "http://localhost:11434/v1",
    "lmstudio":    "http://localhost:1234/v1",
    "openrouter":  "https://openrouter.ai/api/v1",
    "deepseek":    "https://api.deepseek.com/v1",
    "perplexity":  "https://api.perplexity.ai",
    "azure":       "",   # requires AI_BASE_URL — Azure endpoint varies per deployment
    "mistral":     "https://api.mistral.ai/v1",  # Mistral also exposes OpenAI-compat
    "cohere":      "https://api.cohere.ai/compatibility/v1",  # Cohere OpenAI-compat layer
}

# Anthropic /v1/messages compatible — no extra package beyond `anthropic`.
_ANTHROPIC_COMPAT_URLS = {
    "anthropic": "https://api.anthropic.com",
    "claude":    "https://api.anthropic.com",
}


def _openai_compat_call(api_key: str, base_url: str, model: str,
                         system_prompt: str, user_prompt: str,
                         max_tokens: int, skip_ssl: bool,
                         extra_headers: dict = None) -> str:
    """
    Raw httpx call to any OpenAI /v1/chat/completions endpoint.
    No openai SDK required — uses only httpx which is a core dependency.
    """
    import httpx, json as _json

    url = base_url.rstrip("/") + "/chat/completions"
    headers = {
        "Content-Type":  "application/json",
        "Authorization": f"Bearer {api_key or 'not-required'}",
    }
    if extra_headers:
        headers.update(extra_headers)

    payload = {
        "model":      model,
        "max_tokens": max_tokens,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
    }

    _timeout = httpx.Timeout(connect=30.0, read=None, write=60.0, pool=10.0)
    with httpx.Client(verify=not skip_ssl, timeout=_timeout) as http:
        resp = http.post(url, headers=headers, json=payload)

    if resp.status_code == 401:
        raise ValueError(f"401 Unauthorized — check AI_API_KEY in .env\nEndpoint: {url}")
    if resp.status_code == 403:
        raise ValueError(f"403 Forbidden — your key may lack access to model '{model}'\nEndpoint: {url}")
    if resp.status_code not in (200, 201):
        try:
            err = resp.json().get("error", {})
            msg = err.get("message", resp.text[:400]) if isinstance(err, dict) else str(err)
        except Exception:
            msg = resp.text[:400]
        raise ValueError(f"HTTP {resp.status_code} from {url}:\n{msg}")

    data = resp.json()
    return data["choices"][0]["message"]["content"]


def _anthropic_compat_call(api_key: str, base_url: str, model: str,
                            system_prompt: str, user_prompt: str,
                            max_tokens: int, skip_ssl: bool) -> str:
    """
    Raw httpx call to any Anthropic /v1/messages endpoint.
    Uses the anthropic SDK (already a core dependency).
    """
    import anthropic, httpx

    _timeout = httpx.Timeout(connect=30.0, read=None, write=60.0, pool=10.0)
    http_client = httpx.Client(verify=not skip_ssl, timeout=_timeout)
    client = anthropic.Anthropic(
        api_key=api_key,
        base_url=base_url or "https://api.anthropic.com",
        http_client=http_client,
    )
    message = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        system=system_prompt,
        messages=[{"role": "user", "content": user_prompt}],
    )
    return message.content[0].text


def _call_provider(provider: str, api_key: str, base_url: str, model: str,
                   skip_ssl: bool, system_prompt: str, user_prompt: str,
                   max_tokens: int = 8000) -> str:
    """
    Route an AI completion request to the configured provider.
    No extra pip installs needed for Anthropic, OpenAI, Grok, Groq,
    Together, Ollama, DeepSeek, OpenRouter, Perplexity, Mistral, Cohere,
    Azure OpenAI, or any other OpenAI-compatible endpoint.

    Parameters
    ----------
    provider      : AI_PROVIDER value from .env (e.g. "anthropic", "grok")
    api_key       : API key / token (AI_API_KEY in .env)
    base_url      : Base URL override — optional for providers with a default
    model         : Model name (AI_MODEL in .env)
    skip_ssl      : Disable TLS cert verification (SKIP_SSL_VERIFY=true)
    system_prompt : System instructions sent to the model
    user_prompt   : The report generation prompt
    max_tokens    : Maximum tokens to generate per call
    """
    p = provider.strip().lower() if provider else "anthropic"

    # ── Anthropic / Claude ────────────────────────────────────────────────────
    # Uses the anthropic SDK (core dependency — already installed).
    if p in _ANTHROPIC_COMPAT_URLS or p == "anthropic":
        _base = base_url or _ANTHROPIC_COMPAT_URLS.get(p, "https://api.anthropic.com")
        return _anthropic_compat_call(api_key, _base, model,
                                      system_prompt, user_prompt,
                                      max_tokens, skip_ssl)

    # ── Google Gemini ─────────────────────────────────────────────────────────
    # Google's SDK is required — their REST API needs an API-key query param
    # and a different request/response shape not worth reimplementing manually.
    # Install: pip3 install google-generativeai --break-system-packages
    elif p == "gemini":
        try:
            import google.generativeai as genai
        except ImportError:
            raise ValueError(
                "AI_PROVIDER=gemini needs one extra package (Google's own SDK):\n"
                "  pip3 install google-generativeai --break-system-packages\n"
                "All other providers work without any extra install."
            )
        genai.configure(api_key=api_key)
        _mdl = genai.GenerativeModel(
            model_name=model or "gemini-1.5-pro",
            system_instruction=system_prompt,
        )
        response = _mdl.generate_content(
            user_prompt,
            generation_config=genai.GenerationConfig(max_output_tokens=max_tokens),
        )
        return response.text

    # ── OpenAI-compatible (ChatGPT, Grok, Groq, Together, Ollama, Mistral,
    #    Cohere, DeepSeek, OpenRouter, Perplexity, Azure, LM Studio, vLLM …)
    # Pure httpx — no extra packages needed.
    elif p in _OPENAI_COMPAT_URLS:
        _base = base_url or _OPENAI_COMPAT_URLS[p]
        if not _base:
            raise ValueError(
                f"AI_PROVIDER={p} requires AI_BASE_URL to be set in .env\n"
                f"(Azure endpoints vary per deployment — there is no single default URL)"
            )
        # OpenRouter expects a site/app identifier in headers
        extra = {"HTTP-Referer": "https://github.com/vajra-ai/vajra",
                 "X-Title": "VAJRA"} if p == "openrouter" else None
        return _openai_compat_call(api_key, _base, model,
                                   system_prompt, user_prompt,
                                   max_tokens, skip_ssl,
                                   extra_headers=extra)

    # ── Unknown provider — try OpenAI-compat with whatever base_url is set ───
    else:
        print(f"[WARN] Unknown AI_PROVIDER='{p}'. "
              f"Attempting OpenAI-compatible call to: {base_url or '(no base URL set)'}")
        if not base_url:
            raise ValueError(
                f"Unknown AI_PROVIDER='{p}' and no AI_BASE_URL set.\n"
                "Either set AI_PROVIDER to a known provider, or set AI_BASE_URL "
                "to your endpoint and AI_PROVIDER=openai (if OpenAI-compatible) "
                "or AI_PROVIDER=anthropic (if Anthropic-compatible)."
            )
        return _openai_compat_call(api_key, base_url, model,
                                   system_prompt, user_prompt,
                                   max_tokens, skip_ssl)


def generate_report(report_type, client, tester, scope, date, findings_text, template_name,
                    target_system="", tested_from="", assessment_team="",
                    engagement_ref=None, classification=None, report_version=None):
    """
    Call the configured AI provider and generate the report.
    Supports any Anthropic-compatible or OpenAI-compatible endpoint.
    Configure via AI_PROVIDER, AI_API_KEY, AI_BASE_URL, AI_MODEL in .env.
    """
    import httpx

    # ── Load config ──────────────────────────────────────────────────────────
    api_key  = os.environ.get("AI_API_KEY", "").strip()
    base_url = os.environ.get("AI_BASE_URL", "").strip()
    model    = os.environ.get("AI_MODEL", "claude-sonnet-4-6").strip()
    provider = os.environ.get("AI_PROVIDER", "anthropic").strip().lower()
    skip_ssl = os.environ.get("SKIP_SSL_VERIFY", "false").strip().lower() == "true"

    if not api_key:
        raise ValueError("AI_API_KEY not set in .env file")
    if not base_url:
        raise ValueError(
            "AI_BASE_URL not set in .env\n"
            "Examples:\n"
            "  Anthropic:  AI_BASE_URL=https://api.anthropic.com\n"
            "  OpenAI:     AI_BASE_URL=https://api.openai.com/v1\n"
            "  Ollama:     AI_BASE_URL=http://localhost:11434/v1\n"
            "  Proxy:      AI_BASE_URL=https://your-gateway.example.com"
        )

    if not base_url.startswith(("http://", "https://")):
        base_url = "https://" + base_url

    base_url = base_url.rstrip("/")

    print(f"[INFO] Provider: {provider}")
    print(f"[INFO] Endpoint: {base_url}")
    print(f"[INFO] Model:    {model}")
    print(f"[INFO] SSL:      {'skip' if skip_ssl else 'verify'}")

    # ── Build prompt ─────────────────────────────────────────────────────────
    knowledge = load_knowledge_base(
        report_type=report_type,
        scope=scope,
        findings_text=findings_text,
        target_system=target_system,
    )
    # Build meta dict for placeholder injection
    _meta_for_tpl = {
        "client":          client or "",
        "tester":          tester or "",
        "date":            date or "",
        "scope":           scope or "",
        "target_system":   target_system or "",
        "tested_from":     tested_from or "",
        "assessment_team": assessment_team or "",
        "engagement_ref":  engagement_ref or f"VAJRA-{datetime.now().year}-001",
        "classification":  classification or "CONFIDENTIAL — NOT FOR DISTRIBUTION",
        "report_version":  report_version or "1.0 — Final",
    }

    _date  = date or datetime.now().strftime("%Y-%m-%d")
    _client = client or "Client Name"
    _tester = tester or "Security Tester"
    _scope  = scope or "As defined in Rules of Engagement"
    _target = target_system or scope or "As defined in Rules of Engagement"
    _from   = tested_from or "External Network"
    _team   = assessment_team or tester or "Security Team"

    print(f"[GEN] Template selected: '{template_name or '(none)'}'")
    print(f"[GEN] Report type: {report_type}")
    print(f"[GEN] KB size: {len(knowledge)} chars")

    # ── KEY FIX: when a template is selected, use it as the output blueprint ──
    # The template replaces the hardcoded section list entirely.
    # When no template → fall back to hardcoded REPORT_PROMPTS structure.
    tpl_path = (TEMPLATES_DIR / template_name) if template_name else None
    tpl_exists = tpl_path and tpl_path.exists()

    if template_name and tpl_exists:
        # Load and inject placeholders into template
        template_structure = build_template_context(template_name, meta=_meta_for_tpl)
        print(f"[GEN] Using SELECTED TEMPLATE — template drives output structure")
        prompt = f"""Generate a professional security assessment report in Markdown.

ENGAGEMENT DATA:
- Client: {_client}
- Lead Assessor: {_tester} | Assessment Team: {_team}
- Date: {_date} | Target: {_target} | Scope: {_scope} | Tested From: {_from}

FINDINGS PROVIDED:
{findings_text or "No findings provided."}

KNOWLEDGE BASE (use for CWE IDs, OWASP categories, severity definitions, remediation):
{knowledge}

{template_structure}

INSTRUCTIONS:
1. Follow the template structure EXACTLY — do not add or remove sections.
2. Fill every [AI: ...] placeholder with real professional security content.
3. Replace any remaining {{{{placeholder}}}} tokens with appropriate data.
4. For each finding use: ID F-001/F-002/F-003, Severity, CVSS v3.1, CWE, CVE, OWASP, Steps to Reproduce, Evidence, Impact, Remediation.
5. Use knowledge base for accurate CWE IDs, OWASP mapping, and remediation guidance.
6. Calculate remediation deadlines from {_date}: Critical+2d, High+7d, Medium+30d, Low+90d.
7. Output ONLY clean Markdown — no preamble, no commentary."""

    else:
        # No template selected — use hardcoded default structure
        if template_name and not tpl_exists:
            print(f"[GEN] WARNING: Template '{template_name}' not found — using default structure")
        else:
            print(f"[GEN] No template — using default structure")
        template_structure = "No template selected — use standard professional report structure."
        prompt = REPORT_PROMPTS.get(report_type, REPORT_PROMPTS["pentest"]).format(
            knowledge=knowledge,
            template_structure=template_structure,
            client=_client,
            tester=_tester,
            date=_date,
            scope=_scope,
            findings=findings_text or "No findings provided.",
            target_system=_target,
            tested_from=_from,
            assessment_team=_team,
        )

    try:
        return _call_provider(provider, api_key, base_url, model, skip_ssl,
                              SYSTEM_PROMPT, prompt, max_tokens=16000)
    except ValueError:
        raise
    except httpx.ConnectError as e:
        msg = str(e).lower()
        if "ssl" in msg or "certificate" in msg or "tls" in msg:
            raise ValueError(
                f"SSL/TLS error connecting to {base_url}\n"
                f"Add SKIP_SSL_VERIFY=true to your .env file.\nDetail: {e}"
            )
        raise ValueError(
            f"Cannot connect to {base_url}\n"
            f"Check AI_BASE_URL in .env and network access.\nDetail: {e}"
        )
    except httpx.ReadTimeout:
        raise ValueError(
            "The AI provider is taking too long to respond.\n"
            "The API may be under load or the request is too large.\n"
            "Try again in a moment, or reduce the amount of findings/notes."
        )
    except httpx.ConnectTimeout:
        raise ValueError(
            f"Connection to {base_url} timed out.\n"
            "Check AI_BASE_URL in .env and your network connection."
        )
    except httpx.TransportError as e:
        msg = str(e).lower()
        if "ssl" in msg or "certificate" in msg or "tls" in msg:
            raise ValueError(
                f"SSL/TLS error connecting to {base_url}\n"
                f"Add SKIP_SSL_VERIFY=true to your .env file.\nDetail: {e}"
            )
        raise ValueError(
            f"Network error connecting to {base_url}\n"
            f"Check AI_BASE_URL in .env and your network connection.\nDetail: {e}"
        )
    except Exception as e:
        msg = str(e).lower()
        if "ssl" in msg or "certificate" in msg or "tls" in msg:
            raise ValueError(
                f"SSL/TLS error connecting to {base_url}\n"
                f"Add SKIP_SSL_VERIFY=true to your .env file.\nDetail: {e}"
            )
        traceback.print_exc()
        raise ValueError(f"{type(e).__name__}: {str(e)}")


# ── Chunked Generation Engine ─────────────────────────────────────────────────
# For large assessments (10+ findings), splits generation into:
#   Chunk 0: Report skeleton (exec summary, engagement overview, summary table)
#   Chunks 1..N: Batches of FINDINGS_PER_CHUNK findings each (parallel AI calls)
# Assembled into a single markdown then passed to DOCX builder.

FINDINGS_PER_CHUNK = 3   # findings per AI call — 3 per batch reduces 504 gateway timeout risk
_progress_store = {}     # job_id -> { status, pct, message, result, error }
_progress_lock = threading.Lock()

def _set_progress(job_id, pct, message, status="running", result=None, error=None):
    with _progress_lock:
        _progress_store[job_id] = {
            "status": status, "pct": pct, "message": message,
            "result": result, "error": error,
            "updated": datetime.now().isoformat()
        }

def _parse_findings_input(findings_text):
    """
    Parse the user-supplied findings_text into a list of individual finding dicts.
    Supports:
      - Structured text: 'Title: X\\nSeverity: Y\\nDescription: Z'
      - Bullet list: '- SQL Injection | Critical | ...'
      - Numbered list: '1. SQL Injection | Critical'
      - Plain prose per finding (one paragraph = one finding)
    Returns list of {'id': 'F-001', 'raw': '<finding text>'}
    """
    import re
    text = findings_text.strip()
    findings = []

    # Already has ### F-NNN headings (pre-structured)
    if re.search(r'^#{1,4}\s+[A-Z]-\d{3}', text, re.MULTILINE):
        blocks = re.split(r'(?=^#{1,4}\s+[A-Z]-\d{3})', text, flags=re.MULTILINE)
        for b in blocks:
            b = b.strip()
            if b and re.match(r'^#{1,4}\s+[A-Z]-\d{3}', b):
                findings.append(b)
        return findings

    # Numbered or bulleted list items
    items = re.split(r'\n(?=\d+[\.\)]\s+|\-\s+|\*\s+)', text)
    if len(items) > 1:
        for item in items:
            item = item.strip().lstrip('0123456789.-) *').strip()
            if len(item) > 10:
                findings.append(item)
        return findings

    # Double-newline separated paragraphs
    paras = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip() and len(p.strip()) > 20]
    if len(paras) > 1:
        return paras

    # Single block — treat as one finding
    return [text] if text else []

def _call_ai_for_chunk(chunk_type, chunk_data, engagement_ctx, knowledge, api_cfg):
    """
    Make a single AI API call for one chunk.
    chunk_type: 'skeleton' | 'findings_batch'
    chunk_data: for skeleton=None, for findings_batch={'findings': [...], 'start_idx': N, 'total': M}
    Returns markdown string.
    """
    import httpx

    api_key  = api_cfg["api_key"]
    base_url = api_cfg["base_url"]
    model    = api_cfg["model"]
    skip_ssl = api_cfg["skip_ssl"]
    provider = api_cfg.get("provider", "anthropic")

    ctx = engagement_ctx
    date = ctx.get("date", "")
    client = ctx.get("client", "Client")
    tester = ctx.get("tester", "Security Tester")
    report_type = ctx.get("report_type", "pentest")
    total_findings = ctx.get("total_findings", 0)

    if chunk_type == "skeleton":
        all_findings_summary = ctx.get("all_findings_summary", "")
        tpl_structure = ctx.get("template_structure", "")
        _rtype = report_type.lower() if report_type else "pentest"

        # ── PHISHING skeleton ─────────────────────────────────────────────────
        if _rtype == "phishing":
            prompt = f"""Generate the opening sections of a professional phishing simulation report in Markdown.
This is a SIMULATED phishing awareness exercise — NOT an offensive attack report.
Write in a constructive, educational tone focused on user behaviour metrics and improvement.

ENGAGEMENT:
- Client: {client}
- Lead Assessor: {tester} | Team: {ctx.get('assessment_team', tester)}
- Date: {date} | Target Audience: {ctx.get('target_system', '')} | Scope: {ctx.get('scope', '')}
- Total Observations / Campaign Events: {total_findings}

CAMPAIGN OBSERVATIONS:
{all_findings_summary}

KNOWLEDGE BASE:
{knowledge}

OUTPUT THESE SECTIONS ONLY (use exact headings for template mapping):

# Executive Summary
[2-3 paragraphs. Board-level. State: campaign objective, total recipients, overall click/submission/report rates,
risk rating, key behavioural finding, recommended actions. Reference {client} throughout.
Tone: constructive — focus on awareness improvement, not blame.]

# Results
[2-3 sentences summarising overall campaign results: total recipients, click rate, credential submission rate, report rate, overall risk assessment.]

## Campaign mapped to Groups
[REQUIRED — output exactly this table format:]
| # | Campaign Name | Groups | Total Members | Phished |
|---|---|---|---|---|
[One row per campaign from the observations. Use the campaign names/groups from the findings.]

## Detailed Statistics and Metrics
[One paragraph summarising the most notable per-campaign findings and user behaviour patterns.]

# Recommendation
[2-3 actionable recommendations specific to {client}'s results. Focus on: training, policy, technical controls.]

[Second paragraph: specific target improvement goal — e.g. "Aim to reduce click rate below X% within Y months"]

## Remediation
[Specific remediation steps: mandatory training for clickers, DMARC/SPF/DKIM controls, MFA improvements,
awareness programme cadence. Keep to 3-5 practical steps. NO offensive security terminology.]

## Phishing Resilience Score
[Calculate: Non-Click Rate = 100% - click_rate%. Score = (Non-Click Rate × 0.4) + (Report Rate × 0.6).
Show the calculation clearly. Interpret as: 0-30% High Risk, 31-50% Moderate, 51-70% Low Risk, 71-100% Excellent.]

OUTPUT ONLY the sections above. Do NOT include vulnerability findings, CVSS scores, or CWE IDs."""

        # ── RED TEAM skeleton ─────────────────────────────────────────────────
        elif _rtype == "redteam":
            prompt = f"""Generate the opening sections of a professional red team engagement report in Markdown.

ENGAGEMENT:
- Client: {client}
- Lead Assessor: {tester} | Team: {ctx.get('assessment_team', tester)}
- Date: {date} | Target: {ctx.get('target_system', '')} | Objective: {ctx.get('scope', '')}
- Total TTP Observations: {total_findings}

ATTACK CHAIN SUMMARY:
{all_findings_summary}

KNOWLEDGE BASE:
{knowledge}

OUTPUT THESE SECTIONS ONLY:

# Executive Summary
[2-3 paragraphs. State objective, overall outcome (crown jewel reached or not), key attack paths,
detection rate, dwell time. Business risk framing. Reference {client} throughout.]

## Overall Risk Rating
[One sentence with justification based on attack success and detection gaps.]

## Key Findings
[4-5 bullet points — most critical gaps exposed]

# 1. Engagement Overview

## 1.1 Objective and Rules of Engagement
[Objective, scope, assumed breach or full kill-chain, constraints]

## 1.2 Methodology
[MITRE ATT&CK framework phases used, tools, TTP categories]

## 1.3 Limitations
[Time-boxed, non-destructive constraints, scope limitations]

# 2. Attack Path Summary
[High-level kill chain: Initial Access → Execution → Persistence → Lateral Movement → Exfiltration]

## 2.1 Timeline Overview
[Key events table: Date/Time | Phase | Action | Outcome | Detected Y/N]

OUTPUT ONLY the sections above. Do NOT output individual TTP finding blocks."""

        # ── PENTEST skeleton (default) ────────────────────────────────────────
        else:
            prompt = f"""Generate the opening sections of a professional penetration test report in Markdown.

ENGAGEMENT:
- Client: {client}
- Lead Assessor: {tester} | Team: {ctx.get('assessment_team', tester)}
- Date: {date} | Target: {ctx.get('target_system', '')} | Scope: {ctx.get('scope', '')}
- Tested From: {ctx.get('tested_from', '')}
- Total Findings: {total_findings}

ALL FINDINGS SUMMARY (for executive summary and table):
{all_findings_summary}

KNOWLEDGE BASE:
{knowledge}

OUTPUT THESE SECTIONS ONLY — do not include detailed finding blocks:

# Executive Summary
[2-3 paragraphs. Board-level. State what was tested, total findings by severity, top risk, business impact,
most critical action. Reference {client} throughout.]

## Overall Risk Rating
[One sentence overall risk with justification.]

## Key Risk Indicators
[4-5 bullet points — most important risks and business impacts]

# 1. Engagement Overview

## 1.1 Scope and Rules of Engagement
[In-scope, out-of-scope, constraints — reference the scope: {ctx.get('scope','')}]

## 1.2 Methodology
[OWASP WSTG v4.2 / PTES / NIST SP 800-115 testing phases]

## 1.3 Risk Rating Methodology
| Severity | CVSS v3.1 | SLA | Definition |
|---|---|---|---|
| Critical | 9.0–10.0 | 24–48 hours | Trivial exploit, full compromise |
| High | 7.0–8.9 | 7 days | Significant breach risk |
| Medium | 4.0–6.9 | 30 days | Requires conditions |
| Low | 0.1–3.9 | 90 days | Minimal standalone impact |
| Info | 0.0 | Best effort | Hardening opportunity |

## 1.4 Limitations
[4 bullet points — time-boxed, point-in-time, non-destructive, scope-limited]

# 2. Findings Summary

## 2.1 Vulnerability Summary Table
| ID | Title | Severity | CVSS | Affected Asset | CWE | Retest |
|---|---|---|---|---|---|---|
[One row per finding. IDs F-001 through F-{str(total_findings).zfill(3)}]

## 2.2 Severity Distribution
| Severity | Count | % of Total |
|---|---|---|
[One row per severity level with correct counts]

OUTPUT ONLY the sections above. Do NOT output any ### F-NNN finding blocks."""

        # Append template structure if selected
        if tpl_structure:
            prompt += f"\n\nSELECTED TEMPLATE — follow this structure:\n{tpl_structure[:3000]}"

    else:  # findings_batch
        batch = chunk_data["findings"]
        start_idx = chunk_data["start_idx"]  # 0-based
        total = chunk_data["total"]
        _rtype = report_type.lower() if report_type else "pentest"

        findings_block = "\n\n".join(
            f"Observation {start_idx+i+1}:\n{f}" for i, f in enumerate(batch)
        )

        # ── PHISHING batch ────────────────────────────────────────────────────
        if _rtype == "phishing":
            prompt = f"""Document the following phishing campaign observations.
This is a SIMULATED phishing awareness exercise. Write constructively — educational tone, no offensive language.

ENGAGEMENT: {client} | {date} | Observations {start_idx+1}–{start_idx+len(batch)} of {total}

CAMPAIGN OBSERVATIONS:
{findings_block}

KNOWLEDGE BASE:
{knowledge}

For EACH observation above, output this format:

## Campaign [N] — [Campaign Name or Pretext]

**Group Targeted:** [Department / User Group]
**Total Recipients:** [N]
**Emails Sent:** [N]
**Clicked:** [N] ([X]%)
**Credentials Submitted:** [N] ([X]%)
**Reported to IT:** [N] ([X]%)
**Start Date:** [date] | **End Date:** [date]
**Risk Level:** High / Medium / Low (based on click rate)

### Observations
[2-3 sentences: what happened, which group was most susceptible, any notable behaviour]

### Email Pretext Used
[Brief description of the phishing pretext — subject line, sender persona, lure type]

### Key Takeaway
[One sentence: what this campaign revealed about {client}'s awareness posture]

[Repeat for every observation. Output ALL {len(batch)} campaign blocks. Do NOT use CVSS, CWE, or vulnerability language.]"""

        # ── RED TEAM batch ────────────────────────────────────────────────────
        elif _rtype == "redteam":
            prompt = f"""Document the following red team TTP observations.

ENGAGEMENT: {client} | {date} | TTPs {start_idx+1}–{start_idx+len(batch)} of {total}

TTP OBSERVATIONS:
{findings_block}

KNOWLEDGE BASE:
{knowledge}

For EACH observation, output this format:

### TTP-[NNN] — [Technique Name] | MITRE: [ATT&CK ID]

**Phase:** [Kill Chain Phase]
**MITRE ATT&CK:** [Txxxx.xxx — Technique Name]
**Tool/Method:** [tool or technique used]
**Target:** [system or asset]
**Outcome:** [Success/Partial/Blocked]
**Detected:** Yes / No | **Detection Method:** [if detected]
**Dwell Time:** [if applicable]

#### Description
[What was done and why it was significant]

#### Evidence
[Command or action taken, result observed]

#### Detection Gap
[Why the blue team did/didn't detect this]

#### Recommendation
[Specific detection or prevention control]

[Repeat for ALL {len(batch)} TTP blocks.]"""

        # ── PENTEST batch (default) ───────────────────────────────────────────
        else:
            finding_lines = []
            for i, f in enumerate(batch):
                fnum = start_idx + i + 1
                finding_lines.append(f"Finding F-{str(fnum).zfill(3)}:\n{f}")
            findings_block = "\n\n".join(finding_lines)

            prompt = f"""Generate ONLY the detailed finding blocks below. No other sections.

ENGAGEMENT CONTEXT:
- Client: {client} | Date: {date} | Report Type: {report_type}
- Findings F-{str(start_idx+1).zfill(3)} through F-{str(start_idx+len(batch)).zfill(3)} of {total} total.

FINDINGS TO DOCUMENT:
{findings_block}

KNOWLEDGE BASE (use for CWE, OWASP, remediation):
{knowledge}

For EACH finding output EXACTLY this format:

### F-[NNN] — [Title] | [Severity] | CVSS: [score]

**Severity:** [Critical/High/Medium/Low/Info]
**CVSS v3.1 Score:** [score]
**CVSS Vector:** `AV:X/AC:X/PR:X/UI:X/S:X/C:X/I:X/A:X`
**CWE:** CWE-[ID] — [Name]
**CVE:** [CVE-YYYY-NNNNN or No CVE]
**OWASP:** A0X:2021 — [Category]
**Affected Host:** [URL or hostname]
**Root Cause:** [one sentence]
**Remediation Deadline:** [Critical+2d, High+7d, Medium+30d, Low+90d from {date}]
**Retest Status:** Pending Retest

#### Description
[2 paragraphs]

#### Steps to Reproduce
1. [step] 2. [step] 3. [confirm]

#### Evidence
```[request/command]``` ```[response/output]```

#### Business Impact
- **[Category]:** [impact for {client}]

#### Remediation
1. **Immediate:** [fix] 2. **Defense-in-depth:** [control] 3. **Verification:** [check]

[Output ALL {len(batch)} findings. Do NOT skip any.]"""

    return _call_provider(provider, api_key, base_url, model, skip_ssl,
                          SYSTEM_PROMPT, prompt, max_tokens=8000)


def generate_report_chunked(job_id, report_type, client, tester, scope, date,
                             findings_text, template_name,
                             target_system="", tested_from="", assessment_team="",
                             engagement_ref=None, classification=None, report_version=None,
                             store_cb=None):
    """
    Chunked generation for large assessments.
    Runs in a background thread, updates _progress_store[job_id].
    store_cb(content, meta) called on completion to save to DB.
    """
    try:
        _set_progress(job_id, 2, "Parsing findings…")

        # ── Load API config ──
        api_cfg = {
            "api_key":  os.environ.get("AI_API_KEY", "").strip(),
            "base_url": os.environ.get("AI_BASE_URL", "").strip(),
            "model":    os.environ.get("AI_MODEL", "claude-sonnet-4-6").strip(),
            "skip_ssl": os.environ.get("SKIP_SSL_VERIFY", "false").lower() == "true",
            "provider": os.environ.get("AI_PROVIDER", "anthropic").strip().lower(),
        }
        if not api_cfg["base_url"].startswith(("http://", "https://")):
            api_cfg["base_url"] = "https://" + api_cfg["base_url"]
        api_cfg["base_url"] = api_cfg["base_url"].rstrip("/")

        knowledge = load_knowledge_base(
            report_type=report_type,
            scope=scope,
            findings_text=findings_text,
            target_system=target_system,
        )

        # ── Parse findings into individual items ──
        parsed = _parse_findings_input(findings_text)
        total = len(parsed)
        print(f"[CHUNK] Job {job_id}: {total} findings parsed")

        # For small reports (<= FINDINGS_PER_CHUNK), use original single-call path
        if total <= FINDINGS_PER_CHUNK:
            _set_progress(job_id, 5, f"Generating report ({total} findings, single pass)…")
            meta = {
                "report_type": report_type, "client": client, "tester": tester,
                "scope": scope, "date": date, "target_system": target_system,
                "tested_from": tested_from, "assessment_team": assessment_team,
                "engagement_ref": engagement_ref, "classification": classification,
                "report_version": report_version,
            }
            content = generate_report(report_type, client, tester, scope, date,
                                       findings_text, template_name,
                                       target_system=target_system, tested_from=tested_from,
                                       assessment_team=assessment_team,
                                       engagement_ref=engagement_ref,
                                       classification=classification,
                                       report_version=report_version)
            _set_progress(job_id, 95, "Finalising…")
            _chunk_rid = None
            if store_cb:
                _chunk_rid = store_cb(content, meta)
            with _progress_lock:
                _stored_rid = _progress_store.get(job_id, {}).get("result", {}) or {}
                if isinstance(_stored_rid, dict):
                    _chunk_rid = _chunk_rid or _stored_rid.get("report_id")
            _set_progress(job_id, 100, "Complete", status="done", result={"content": content, "report_id": _chunk_rid})
            return

        # ── Large report: chunked parallel generation ──
        # Build a compact summary of all findings for the skeleton call
        summary_lines = []
        for i, f in enumerate(parsed):
            fid = f"F-{str(i+1).zfill(3)}"
            # Extract title/severity from first line of each finding
            first = f.split('\n')[0].strip().lstrip('#').strip()
            summary_lines.append(f"  {fid}: {first[:120]}")
        all_findings_summary = "\n".join(summary_lines)

        # Build template context for chunked path
        _tpl_meta_c = {"client": client, "tester": tester, "date": date, "scope": scope,
                       "target_system": target_system, "tested_from": tested_from,
                       "assessment_team": assessment_team or tester}
        _tpl_path_c = (TEMPLATES_DIR / template_name) if template_name else None
        _tpl_exists_c = bool(_tpl_path_c and _tpl_path_c.exists())
        _tpl_structure_c = build_template_context(template_name, meta=_tpl_meta_c) if _tpl_exists_c else ""

        engagement_ctx = {
            "report_type": report_type, "client": client, "tester": tester,
            "assessment_team": assessment_team or tester,
            "scope": scope, "date": date,
            "target_system": target_system, "tested_from": tested_from,
            "total_findings": total,
            "all_findings_summary": all_findings_summary,
            "template_name": template_name or "",
            "template_structure": _tpl_structure_c,
        }

        # Split findings into batches of FINDINGS_PER_CHUNK
        batches = []
        for start in range(0, total, FINDINGS_PER_CHUNK):
            batch = parsed[start:start + FINDINGS_PER_CHUNK]
            batches.append({"findings": batch, "start_idx": start, "total": total})

        n_chunks = 1 + len(batches)   # 1 skeleton + N finding batches
        print(f"[CHUNK] Job {job_id}: {n_chunks} chunks ({len(batches)} finding batches)")

        results = {}   # chunk_key -> markdown text

        # ── Chunk 0: skeleton (sequential — must run first) ──
        _set_progress(job_id, 8,
            f"Generating report skeleton (exec summary, overview, {total}-finding table)…")
        skeleton_md = _call_ai_for_chunk("skeleton", None, engagement_ctx, knowledge, api_cfg)
        results["skeleton"] = skeleton_md
        print(f"[CHUNK] Job {job_id}: skeleton done ({len(skeleton_md)} chars)")

        # ── Chunks 1..N: finding batches in PARALLEL ──
        completed = 0
        lock = threading.Lock()

        def run_batch(i, batch_data):
            nonlocal completed
            batch_md = _call_ai_for_chunk("findings_batch", batch_data, engagement_ctx, knowledge, api_cfg)
            with lock:
                completed += 1
                pct = 20 + int(completed / len(batches) * 70)
                start_n = batch_data["start_idx"] + 1
                end_n   = batch_data["start_idx"] + len(batch_data["findings"])
                _set_progress(job_id, pct,
                    f"Generated findings F-{str(start_n).zfill(3)}–F-{str(end_n).zfill(3)} "
                    f"({completed}/{len(batches)} batches done)…")
            print(f"[CHUNK] Job {job_id}: batch {i} done ({len(batch_md)} chars)")
            return i, batch_md

        max_workers = min(len(batches), 6)   # max 6 parallel AI calls
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as ex:
            futures = {ex.submit(run_batch, i, b): i for i, b in enumerate(batches)}
            for fut in concurrent.futures.as_completed(futures):
                i, batch_md = fut.result()   # raises on exception
                results[f"batch_{i}"] = batch_md

        # ── Assemble ──
        _set_progress(job_id, 93, "Assembling all sections…")
        parts = [results["skeleton"].strip()]
        for i in range(len(batches)):
            parts.append(results[f"batch_{i}"].strip())

        # Add Remediation Roadmap + Appendix (quick, no AI needed — DOCX builder handles these)
        parts.append("\n# 5. Remediation Roadmap\n\n*Generated by VAJRA — see Remediation Roadmap section.*\n")
        parts.append("\n# 6. Appendix\n\n*Generated by VAJRA — see Appendix section.*\n")

        content = "\n\n".join(parts)

        _set_progress(job_id, 96, "Saving report…")
        meta = {
            "report_type": report_type, "client": client, "tester": tester,
            "scope": scope, "date": date, "target_system": target_system,
            "tested_from": tested_from, "assessment_team": assessment_team,
            "engagement_ref": engagement_ref, "classification": classification,
            "report_version": report_version,
        }
        _large_rid = None
        if store_cb:
            _large_rid = store_cb(content, meta)
        with _progress_lock:
            _stored = _progress_store.get(job_id, {}).get("result", {}) or {}
            if isinstance(_stored, dict):
                _large_rid = _large_rid or _stored.get("report_id")
        _set_progress(job_id, 100, f"Done — {total} findings, {n_chunks} AI calls",
                      status="done", result={"content": content, "report_id": _large_rid})
        print(f"[CHUNK] Job {job_id}: complete. Total content: {len(content)} chars")

    except Exception as e:
        traceback.print_exc()
        _set_progress(job_id, 0, f"Error: {e}", status="error", error=str(e))


# ── DOCX Export ───────────────────────────────────────────────────────────────

def markdown_to_docx(md_content, output_path, template_path=None, title="Security Report", meta=None):
    """
    Route to the right DOCX builder based on template type:
    - .docx template  → docx_template_filler (fills actual template file)
    - .md / no template → generate.js (VAJRA professional styled output)
    meta = { client, tester, date, scope, report_type, ... }
    """
    import subprocess, json, tempfile, os

    # ── Scan logos FIRST so they're available to both paths ──────────────────
    comp_logo = None
    test_logo = None
    try:
        for f in LOGOS_DIR.iterdir():
            if f.suffix.lower() not in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"]:
                continue
            if f.stat().st_size < 100:
                continue
            if f.stem.startswith("company_logo"):
                comp_logo = str(f)
                print(f"[LOGO] Company logo: {f.name} ({f.stat().st_size} bytes)")
            elif f.stem.startswith("tester_logo"):
                test_logo = str(f)
                print(f"[LOGO] Tester logo:  {f.name} ({f.stat().st_size} bytes)")
    except Exception as e:
        print(f"[LOGO] Could not scan logos dir: {e}")

    # ── Custom .docx template → use template filler ───────────────────────────
    if template_path and str(template_path).lower().endswith(".docx"):
        print(f"[DOCX] Using custom template: {Path(template_path).name}")
        try:
            sys.path.insert(0, str(BASE_DIR / "docx_builder"))
            import importlib
            if "docx_template_filler" in sys.modules:
                importlib.reload(sys.modules["docx_template_filler"])
            from docx_template_filler import fill_docx_template
            _meta = dict(meta or {})
            _meta["_company_logo"] = comp_logo
            _meta["_tester_logo"]  = test_logo
            ok = fill_docx_template(
                template_path=template_path,
                output_path=output_path,
                meta=_meta,
                ai_content_text=md_content,
            )
            if ok and Path(output_path).exists():
                print(f"[DOCX] ✓ Custom template filled: {Path(template_path).name}")
                return
            print(f"[WARN] Template filler failed — falling back to generate.js")
        except Exception as e:
            import traceback
            print(f"[WARN] Template filler error: {e}")
            traceback.print_exc()
            print("[WARN] Falling back to generate.js")

    # ── Default: VAJRA styled output via generate.js ──────────────────────────
    builder = BASE_DIR / "docx_builder" / "generate.js"
    if not builder.exists():
        _write_plain_docx(md_content, output_path)
        return

    payload = {
        "content":     md_content,
        "meta":        meta or {},
        "companyLogo": comp_logo,
        "testerLogo":  test_logo,
    }

    with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False, encoding="utf-8") as f:
        json.dump(payload, f)
        tmp = f.name

    try:
        r = subprocess.run(
            ["node", str(builder), tmp, str(output_path)],
            capture_output=True, text=True, timeout=60
        )
        if r.returncode == 0 and Path(output_path).exists():
            return
        print(f"[WARN] Node docx builder failed:\n{r.stderr[:400]}")
    finally:
        try: os.unlink(tmp)
        except: pass

    # Last-resort plain fallback
    _write_plain_docx(md_content, output_path)


def _write_plain_docx(md_content, output_path):
    """Minimal python-docx emergency fallback."""
    from docx import Document
    from docx.shared import RGBColor
    doc = Document()
    sev = {"critical": RGBColor(0xC0,0,0), "high": RGBColor(0xFF,0,0),
           "medium": RGBColor(0xFF,0x92,0), "low": RGBColor(0xFF,0xC0,0)}
    for line in md_content.split("\n"):
        s = line.strip()
        if not s or s == "---": continue
        if s.startswith("# "):   doc.add_heading(s[2:],  1); continue
        if s.startswith("## "):  doc.add_heading(s[3:],  2); continue
        if s.startswith("### "): doc.add_heading(s[4:],  3); continue
        if s.startswith(("- ","* ")):
            doc.add_paragraph().add_run(f"\u2022  {s[2:]}"); continue
        p = doc.add_paragraph()
        c = next((v for k,v in sev.items() if f"severity: {k}" in s.lower()), None)
        run = p.add_run(s.replace("**",""))
        if c: run.font.color.rgb = c; run.bold = True
    doc.save(str(output_path))

# ── Auth helpers ─────────────────────────────────────────────────────────────

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "username" not in session:
            if request.path.startswith("/api/"):
                return jsonify({"error": "Not authenticated", "redirect": "/login"}), 401
            return redirect(url_for("login_page"))
        return f(*args, **kwargs)
    return decorated

def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if session.get("role") != "admin":
            if request.path.startswith("/api/"):
                return jsonify({"error": "Admin access required"}), 403
            return redirect(url_for("index"))
        return f(*args, **kwargs)
    return decorated


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route("/login")
def login_page():
    if "username" in session:
        return redirect(url_for("index"))
    return render_template("login.html")

@app.route("/api/login", methods=["POST"])
def api_login():
    ip = request.remote_addr or "unknown"
    if _check_rate_limit(ip):
        return jsonify({"error": "Too many failed attempts. Try again in 5 minutes."}), 429
    data = request.json or {}
    username = (data.get("username") or "").strip().lower()
    password = data.get("password") or ""
    if not username or not password:
        return jsonify({"error": "Username and password required"}), 400
    # Also rate-limit per username to prevent distributed brute force
    if _check_rate_limit(f"user:{username}"):
        return jsonify({"error": "Account temporarily locked. Try again in 5 minutes."}), 429
    with get_db() as conn:
        user = conn.execute("SELECT * FROM users WHERE username=?", (username,)).fetchone()
    if not user or not _verify_password(password, user["password_hash"]):
        _record_failed(ip)
        _record_failed(f"user:{username}")
        return jsonify({"error": "Invalid username or password"}), 401
    if user["is_disabled"]:
        return jsonify({"error": "Your account has been disabled. Contact your administrator."}), 403
    _clear_failed(ip)
    with get_db() as conn:
        conn.execute("UPDATE users SET last_login=datetime('now') WHERE username=?", (username,))
        conn.commit()
    # Session fixation fix — clear old session before setting new one
    session.clear()
    session.permanent = True
    session["username"]  = username
    session["role"]      = user["role"]
    session["full_name"] = user["full_name"] or username
    return jsonify({"success": True, "username": username, "role": user["role"]})

@app.route("/api/logout", methods=["POST"])
def api_logout():
    session.clear()
    return jsonify({"success": True})

@app.route("/api/me")
def api_me():
    if "username" not in session:
        return jsonify({"authenticated": False}), 401
    role = session["role"]
    perms = {
        "can_generate":      role in ("admin", "analyst"),
        "can_upload":        role in ("admin", "analyst"),
        "can_delete_own":    role in ("admin", "analyst"),
        "can_delete_any":    role == "admin",
        "can_manage_users":  role == "admin",
        "can_view_reports":  True,
        "can_export":        True,
    }
    return jsonify({
        "authenticated": True,
        "username":  session["username"],
        "role":      role,
        "full_name": session["full_name"],
        "permissions": perms,
    })

@app.route("/api/users", methods=["GET"])
@login_required
@admin_required
def list_users():
    with get_db() as conn:
        rows = conn.execute(
            "SELECT id,username,role,full_name,created_at,last_login,is_disabled FROM users ORDER BY id"
        ).fetchall()
    return jsonify([dict(r) for r in rows])

@app.route("/api/users", methods=["POST"])
@login_required
@admin_required
def create_user():
    data = request.json or {}
    username  = (data.get("username") or "").strip().lower()
    password  = data.get("password") or ""
    role      = data.get("role", "analyst")
    full_name = data.get("full_name", "")
    if not username or not password:
        return jsonify({"error": "Username and password required"}), 400
    if len(password) < 8:
        return jsonify({"error": "Password must be at least 8 characters"}), 400
    if not re.match(r'^[a-z0-9_.-]+$', username):
        return jsonify({"error": "Username: lowercase letters, numbers, _ . - only"}), 400
    if role not in ("admin", "analyst", "viewer"):
        return jsonify({"error": "Role must be admin/analyst/viewer"}), 400
    try:
        with get_db() as conn:
            conn.execute(
                "INSERT INTO users (username,password_hash,role,full_name) VALUES (?,?,?,?)",
                (username, _hash_password(password), role, full_name)
            )
            conn.commit()
        return jsonify({"success": True, "username": username})
    except Exception as e:
        return jsonify({"error": f"Username already exists"}), 409

@app.route("/api/users/<username>", methods=["DELETE"])
@login_required
@admin_required
def delete_user(username):
    if username == session["username"]:
        return jsonify({"error": "Cannot delete your own account"}), 400
    with get_db() as conn:
        conn.execute("DELETE FROM users WHERE username=?", (username,))
        conn.commit()
    return jsonify({"success": True})

@app.route("/api/users/<username>/disable", methods=["POST"])
@login_required
@admin_required
def disable_user(username):
    if username == session["username"]:
        return jsonify({"error": "Cannot disable your own account"}), 400
    with get_db() as conn:
        u = conn.execute("SELECT id FROM users WHERE username=?", (username,)).fetchone()
        if not u:
            return jsonify({"error": "User not found"}), 404
        conn.execute("UPDATE users SET is_disabled=1 WHERE username=?", (username,))
        conn.commit()
    return jsonify({"success": True, "message": f"{username} has been disabled"})

@app.route("/api/users/<username>/enable", methods=["POST"])
@login_required
@admin_required
def enable_user(username):
    with get_db() as conn:
        u = conn.execute("SELECT id FROM users WHERE username=?", (username,)).fetchone()
        if not u:
            return jsonify({"error": "User not found"}), 404
        conn.execute("UPDATE users SET is_disabled=0 WHERE username=?", (username,))
        conn.commit()
    return jsonify({"success": True, "message": f"{username} has been enabled"})

@app.route("/api/users/<username>/password", methods=["POST"])
@login_required
def change_password(username):
    if session["role"] != "admin" and session["username"] != username:
        return jsonify({"error": "Not authorised"}), 403
    data = request.json or {}
    new_pass = data.get("password") or ""
    if len(new_pass) < 8:
        return jsonify({"error": "Password must be at least 8 characters"}), 400
    with get_db() as conn:
        conn.execute("UPDATE users SET password_hash=? WHERE username=?",
                     (_hash_password(new_pass), username))
        conn.commit()
    # Invalidate current session if user changed their own password
    if session.get("username") == username:
        session.clear()
    return jsonify({"success": True, "relogin": session.get("username") != username})


@app.route("/")
@login_required
def index():
    with get_db() as conn:
        reports = conn.execute("SELECT * FROM reports ORDER BY created_at DESC LIMIT 50").fetchall()
    templates = [f.name for f in TEMPLATES_DIR.iterdir() if f.suffix in [".docx", ".md", ".html", ".txt"]]
    kb_files = [f.name for f in KNOWLEDGE_DIR.iterdir() if f.suffix in [".md", ".txt"]]
    return render_template("index.html",
                           reports=[dict(r) for r in reports],
                           templates=templates,
                           kb_files=kb_files)


@app.route("/api/status")
@login_required
def api_status():
    api_key  = os.environ.get("AI_API_KEY", "")
    base_url = os.environ.get("AI_BASE_URL", "").strip()
    model    = os.environ.get("AI_MODEL", "claude-sonnet-4-6")
    provider = os.environ.get("AI_PROVIDER", "anthropic")
    skip_ssl = os.environ.get("SKIP_SSL_VERIFY", "false")
    kb   = len(glob.glob(str(KNOWLEDGE_DIR / "*.md"))) + len(glob.glob(str(KNOWLEDGE_DIR / "*.txt")))
    tmpl = len([f for f in TEMPLATES_DIR.iterdir() if f.suffix in [".docx", ".md", ".html"]])
    with get_db() as conn:
        rpts = conn.execute("SELECT COUNT(*) FROM reports").fetchone()[0]
    ready = bool(api_key and base_url)
    # Check logos
    comp_logo_ok = any(
        f.stem.startswith("company_logo") and f.stat().st_size > 100
        for f in LOGOS_DIR.iterdir() if f.suffix.lower() in [".png",".jpg",".jpeg",".webp"]
    ) if LOGOS_DIR.exists() else False
    test_logo_ok = any(
        f.stem.startswith("tester_logo") and f.stat().st_size > 100
        for f in LOGOS_DIR.iterdir() if f.suffix.lower() in [".png",".jpg",".jpeg",".webp"]
    ) if LOGOS_DIR.exists() else False

    return jsonify({
        "api_key_set":  bool(api_key),
        "base_url":     base_url or "⚠ NOT SET",
        "model":        model,
        "provider":     provider,
        "skip_ssl":     skip_ssl,
        "kb_files":     kb,
        "templates":    tmpl,
        "reports":      rpts,
        "ready":        ready,
        "ai_connected": ready,
        "company_logo": comp_logo_ok,
        "tester_logo":  test_logo_ok,
    })


@app.route("/api/parse-tool-import", methods=["POST"])
@login_required
def parse_tool_import():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400
    file = request.files["file"]
    content = file.read().decode("utf-8", errors="ignore")
    try:
        if "NessusClientData" in content or "ReportHost" in content:
            findings = parse_nessus(content)
        elif "<issues>" in content or "<issue>" in content:
            findings = parse_burp(content)
        elif "<nmaprun" in content:
            findings = parse_nmap(content)
        else:
            return jsonify({"error": "Could not detect tool type. Supported: Nessus, Burp, Nmap XML"}), 400
        return jsonify({"findings_text": findings_to_text(findings), "count": len(findings)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/upload-template", methods=["POST"])
@login_required
def upload_template():
    if session.get("role") == "viewer":
        return jsonify({"error": "Viewer role cannot upload files"}), 403
    if "file" not in request.files:
        return jsonify({"error": "No file"}), 400
    file = request.files["file"]
    if file.filename.rsplit(".", 1)[-1].lower() not in ["docx", "md", "html", "txt"]:
        return jsonify({"error": "Unsupported format. Use .docx, .md, or .html"}), 400
    fname = secure_filename(file.filename)
    if not fname:
        return jsonify({"error": "Invalid filename"}), 400
    file.save(str(TEMPLATES_DIR / fname))
    return jsonify({"success": True, "filename": fname})


@app.route("/api/upload-knowledge", methods=["POST"])
@login_required
def upload_knowledge():
    if session.get("role") == "viewer":
        return jsonify({"error": "Viewer role cannot upload files"}), 403
    if "file" not in request.files:
        return jsonify({"error": "No file"}), 400
    file = request.files["file"]
    if file.filename.rsplit(".", 1)[-1].lower() not in ["md", "txt"]:
        return jsonify({"error": "Knowledge files must be .md or .txt"}), 400
    fname = secure_filename(file.filename)
    if not fname:
        return jsonify({"error": "Invalid filename"}), 400
    file.save(str(KNOWLEDGE_DIR / fname))
    return jsonify({"success": True, "filename": fname})

@app.route("/api/upload-logo", methods=["POST"])
@login_required
def upload_logo():
    f = request.files.get("logo")
    if not f: return jsonify({"error": "No file"}), 400
    ext = Path(f.filename).suffix.lower()
    if ext not in [".png", ".jpg", ".jpeg", ".svg", ".webp"]:
        return jsonify({"error": "Invalid type"}), 400
    # save as company_logo or tester_logo based on type param
    logo_type = request.form.get("type", "company")
    logo_type = secure_filename(logo_type)
    fname = f"{logo_type}_logo{ext}"
    f.save(str(LOGOS_DIR / fname))
    return jsonify({"ok": True, "filename": fname, "url": f"/logos/{fname}"})


@app.route("/logos/<path:filename>")
@login_required
def serve_logo(filename):
    # Prevent path traversal — only serve files inside LOGOS_DIR
    from werkzeug.utils import safe_join
    try:
        safe_path = safe_join(str(LOGOS_DIR), secure_filename(filename))
        if not safe_path or not Path(safe_path).exists():
            abort(404)
        return send_file(safe_path)
    except Exception:
        abort(404)


@app.route("/api/logos")
@login_required
def list_logos():
    logos = {}
    for f in LOGOS_DIR.iterdir():
        if f.stem.startswith("company_logo"): logos["company"] = f"/logos/{f.name}"
        if f.stem.startswith("tester_logo"):  logos["tester"]  = f"/logos/{f.name}"
    return jsonify(logos)


@app.route("/api/generate", methods=["POST"])
@login_required
def generate():
    if session.get("role") == "viewer":
        return jsonify({"error": "Viewer role cannot generate reports. Contact admin."}), 403
    # Rate limit generation — max 10 per hour per user
    gen_key = f"gen:{session.get('username','unknown')}"
    if _check_rate_limit(gen_key):
        return jsonify({"error": "Generation rate limit reached. Maximum 10 reports per 5 minutes."}), 429
    _record_failed(gen_key)  # counts towards limit
    data = request.json or {}
    report_type  = data.get("report_type", "pentest")
    client       = data.get("client", "")
    tester       = data.get("tester", "")
    scope        = data.get("scope", "")
    date         = data.get("date", datetime.now().strftime("%Y-%m-%d"))
    findings     = data.get("findings_text", "")
    template     = data.get("template_name", "")
    engagement_ref   = data.get("engagement_ref", "") or ""
    classification   = data.get("classification", "CONFIDENTIAL — NOT FOR DISTRIBUTION") or ""
    report_version   = data.get("report_version", "1.0 — Final") or ""
    _extra           = data.get("extra_fields") or {}
    _reviewer        = _extra.get("reviewer", "") or ""
    _approver        = _extra.get("approver", "") or ""
    _campaign_period = _extra.get("campaign_period", "") or ""
    _platform        = _extra.get("platform", "") or ""
    import json as _json
    _campaigns_raw   = _extra.get("campaigns", [])
    _campaigns_json  = _json.dumps(_campaigns_raw) if _campaigns_raw else ""
    _distribution_list  = _extra.get("distribution_list", "") or ""
    _phished_employees  = _extra.get("phished_employees", "") or ""
    title        = data.get("title", f"{report_type.upper()} — {client or 'Client'} — {date}")

    target_system   = data.get("target_system", "")
    tested_from     = data.get("tested_from", "")
    assessment_team = data.get("assessment_team", "")
    engagement_ref  = data.get("engagement_ref", "") or f"VAJRA-{datetime.now().year}-{str(int(datetime.now().timestamp()))[-4:]}"
    classification  = data.get("classification", "CONFIDENTIAL — NOT FOR DISTRIBUTION")
    report_version  = data.get("report_version", "1.0 — Final")
    extra_fields    = data.get("extra_fields", {}) or {}

    # Merge extra_fields into findings context
    extra_context = ""
    if extra_fields:
        lines = []
        field_labels = {
            "assess_type":"Assessment Type","out_of_scope":"Out-of-Scope",
            "objective":"Primary Objective / Crown Jewels","threat_actor":"Threat Actor Simulated",
            "c2_framework":"C2 Framework","init_access":"Initial Access Vector",
            "crown_reached":"Crown Jewels Reached","detection_rate":"Blue Team Detection Rate",
            "blue_awareness":"Blue Team Awareness","duration":"Engagement Duration",
            "total_recipients":"Total Recipients","campaign_type":"Campaign Type",
            "pretext":"Pretext Used","spoofed_as":"Email Spoofed As",
            "phish_domain":"Phishing Domain","mfa_bypass":"MFA Bypass Result",
            "emails_sent":"Emails Sent","clicked":"Links Clicked",
            "creds":"Credentials Harvested","reported":"Reports from Staff",
        }
        for k, v in extra_fields.items():
            if v and k not in ("target_system", "tested_from", "scope"):
                label = field_labels.get(k, k.replace("_", " ").title())
                lines.append(f"- {label}: {v}")
        if lines:
            extra_context = "\n\nEXTRA ENGAGEMENT DATA:\n" + "\n".join(lines)

    findings_with_extra = findings + extra_context

    if not findings.strip():
        return jsonify({"error": "No findings provided. Please add findings before generating."}), 400

    # ── Decide: chunked (large) or direct (small) ──
    parsed_count = len(_parse_findings_input(findings))
    use_chunked = parsed_count > FINDINGS_PER_CHUNK

    import uuid
    job_id = str(uuid.uuid4())[:8]
    _created_by = session.get("username", "unknown")
    _visibility = data.get("visibility", "team")

    def store_cb(content, meta):
        try:
            with get_db() as conn:
                cur = conn.execute("""
                    INSERT INTO reports (title, report_type, client, tester, scope,
                                         status, findings_raw, generated_content, template_used,
                                         target_system, tested_from, assessment_team, created_by, visibility,
                                         reviewer, approver, campaign_period, platform,
                                         campaigns_json, distribution_list, phished_employees,
                                         engagement_ref, classification, report_version)
                    VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                """, (title, report_type, client, tester, scope, "generated",
                      findings, content, template,
                      target_system, tested_from, assessment_team,
                      _created_by, _visibility,
                      _reviewer, _approver, _campaign_period, _platform,
                      _campaigns_json, _distribution_list, _phished_employees,
                      engagement_ref, classification, report_version))
                conn.commit()
                report_id = cur.lastrowid
            with _progress_lock:
                if job_id in _progress_store:
                    if _progress_store[job_id].get("result") is None:
                        _progress_store[job_id]["result"] = {}
                    _progress_store[job_id]["result"]["report_id"] = report_id
            return report_id  # ← return so caller can use it directly
        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"[STORE] Error saving report: {e}")
            raise

    if use_chunked:
        # Start background thread, return job_id immediately
        _set_progress(job_id, 1, f"Starting chunked generation ({parsed_count} findings, "
                      f"{(parsed_count + FINDINGS_PER_CHUNK - 1) // FINDINGS_PER_CHUNK + 1} AI calls)…")
        t = threading.Thread(
            target=generate_report_chunked,
            args=(job_id, report_type, client, tester, scope, date,
                  findings_with_extra, template),
            kwargs=dict(target_system=target_system, tested_from=tested_from,
                        assessment_team=assessment_team, engagement_ref=engagement_ref,
                        classification=classification, report_version=report_version,
                        store_cb=store_cb),
            daemon=True
        )
        t.start()
        return jsonify({
            "success": True,
            "async": True,
            "job_id": job_id,
            "total_findings": parsed_count,
            "total_chunks": (parsed_count + FINDINGS_PER_CHUNK - 1) // FINDINGS_PER_CHUNK + 1,
            "message": f"Large report ({parsed_count} findings) generating in background. Poll /api/generate/progress/{job_id}"
        })

    else:
        # Small report: direct synchronous generation (original behaviour)
        try:
            _set_progress(job_id, 5, f"Generating ({parsed_count} findings)…")
            content = generate_report(report_type, client, tester, scope, date,
                                      findings_with_extra, template,
                                      target_system=target_system, tested_from=tested_from,
                                      assessment_team=assessment_team,
                                      engagement_ref=engagement_ref,
                                      classification=classification,
                                      report_version=report_version)
            try:
                report_id = store_cb(content, {})  # store_cb now returns report_id directly
                print(f"[GENERATE] Report saved, id={report_id}")
            except Exception as store_err:
                print(f"[GENERATE] store_cb failed: {store_err}")
                report_id = None
            _set_progress(job_id, 100, "Done", status="done", result={"content": content, "report_id": report_id})

            return jsonify({"success": True, "async": False,
                            "job_id": job_id, "report_id": report_id, "content": content})
        except ValueError as e:
            return jsonify({"error": str(e)}), 400
        except Exception as e:
            traceback.print_exc()
            return jsonify({"error": f"{type(e).__name__}: {str(e)}"}), 500


@app.route("/api/generate/progress/<job_id>", methods=["GET"])
@login_required
def generate_progress(job_id):
    """Poll this endpoint to track async chunked generation progress."""
    with _progress_lock:
        state = _progress_store.get(job_id)
    if not state:
        return jsonify({"error": "Job not found"}), 404
    return jsonify(state)


@app.route("/api/reports")
@login_required
def list_reports():
    username = session["username"]
    role     = session["role"]
    with get_db() as conn:
        if role == "admin":
            # Admin sees all reports
            rows = conn.execute(
                "SELECT id,title,report_type,client,tester,created_at,created_by,visibility "
                "FROM reports ORDER BY id DESC"
            ).fetchall()
        else:
            # Non-admin sees a report if ANY of:
            # 1. visibility='public'  (everyone sees it)
            # 2. visibility='team'    (all analysts see it)
            # 3. created_by=username  (own reports, any visibility)
            # 4. explicitly shared with username via report_shares
            # Private reports NOT in shares are NOT visible
            rows = conn.execute(
                "SELECT DISTINCT r.id,r.title,r.report_type,r.client,r.tester,"
                "r.created_at,r.created_by,r.visibility "
                "FROM reports r "
                "LEFT JOIN report_shares s "
                "  ON r.id=s.report_id AND s.shared_with=? "
                "WHERE r.visibility IN ('public','team') "
                "   OR r.created_by=? "
                "   OR s.report_id IS NOT NULL "
                "ORDER BY r.id DESC",
                (username, username)
            ).fetchall()
    return jsonify([dict(r) for r in rows])

@app.route("/api/report/<int:rid>/visibility", methods=["POST"])
@login_required
def set_visibility(rid):
    data = request.json or {}
    vis = data.get("visibility", "team")
    if vis not in ("private", "team", "public"):
        return jsonify({"error": "visibility must be private, team, or public"}), 400
    with get_db() as conn:
        r = conn.execute("SELECT created_by FROM reports WHERE id=?", (rid,)).fetchone()
    if not r:
        return jsonify({"error": "Report not found"}), 404
    if session["role"] != "admin" and r["created_by"] != session["username"]:
        return jsonify({"error": "Only the report creator or admin can change visibility"}), 403
    with get_db() as conn:
        conn.execute("UPDATE reports SET visibility=? WHERE id=?", (vis, rid))
        conn.commit()
    return jsonify({"success": True, "visibility": vis})


@app.route("/api/report/<int:rid>")
@login_required
def get_report(rid):
    with get_db() as conn:
        r = conn.execute("SELECT * FROM reports WHERE id=?", (rid,)).fetchone()
    if not r:
        abort(404)
    r = dict(r)
    if not can_access_report(r, session.get("username"), session.get("role", "viewer")):
        return jsonify({"error": "Access denied"}), 403
    return jsonify(r)


@app.route("/api/export/<int:rid>/<fmt>")
@login_required
def export_report(rid, fmt):
    try:
        with get_db() as conn:
            r = conn.execute("SELECT * FROM reports WHERE id=?", (rid,)).fetchone()
        if not r:
            return jsonify({"error": f"Report ID {rid} not found. Generate a report first."}), 404
        r = dict(r)
        # Check access using full access control (visibility + shares)
        if not can_access_report(r, session.get("username"), session.get("role", "viewer")):
            return jsonify({"error": "Access denied — you do not have permission to access this report"}), 403
        content = r.get("generated_content", "")
        if not content:
            return jsonify({"error": "Report has no content to export."}), 400
        template_name = r.get("template_used", "")
        safe = "".join(c for c in r["title"] if c.isalnum() or c in " -_")[:60].strip()
        if not safe:
            safe = f"report_{rid}"

        # Build meta once — used by both docx and pdf branches
        report_meta = {
            "client":          r.get("client", ""),
            "tester":          r.get("tester", ""),
            "date":            (r.get("created_at", "") or "")[:10],
            "scope":           r.get("scope", ""),
            "report_type":     r.get("report_type", "pentest"),
            "target_system":   r.get("target_system", "") or r.get("scope", ""),
            "tested_from":     r.get("tested_from", "") or "External Network",
            "assessment_team": r.get("assessment_team", "") or r.get("tester", ""),
            # Additional fields for .docx template filler
            "engagement_ref":  r.get("engagement_ref", "") or f"ENG-{(r.get('created_at','2026') or '2026')[:4]}-001",
            "classification":  r.get("classification", "CONFIDENTIAL — NOT FOR DISTRIBUTION"),
            "report_version":  r.get("report_version", "") or "1.0 — Final",
            "title":           r.get("title", "") or "Security Assessment Report",
            "report_id":       str(r.get("id", "")),
            "author":          r.get("tester", ""),
            # Phishing Campaign fields
            "reviewer":         r.get("reviewer", "") or "",
            "approver":         r.get("approver", "") or "",
            "campaign_period":  r.get("campaign_period", "") or "",
            "platform":         r.get("platform", "") or "",
            "distribution_list": r.get("distribution_list", "") or "",
            "phished_employees": r.get("phished_employees", "") or "",
        }
        # Parse campaigns JSON back into structured data for filler
        import json as _json
        _camps_raw = r.get("campaigns_json", "") or ""
        if _camps_raw:
            try:
                report_meta["campaigns"] = _json.loads(_camps_raw)
            except Exception:
                report_meta["campaigns"] = []
        else:
            report_meta["campaigns"] = []

        if fmt == "md":
            p = REPORTS_DIR / f"{safe}.md"
            p.write_text(content, encoding="utf-8")
            return send_file(str(p), as_attachment=True, download_name=f"{safe}.md")

        elif fmt == "docx":
            p = REPORTS_DIR / f"{safe}.docx"
            tp = (TEMPLATES_DIR / template_name) if template_name else None
            tp = str(tp) if tp and tp.exists() and tp.suffix == ".docx" else None
            markdown_to_docx(content, str(p), tp,
                             title=r.get("title","Security Report"),
                             meta=report_meta)
            return send_file(str(p), as_attachment=True, download_name=f"{safe}.docx")

        elif fmt == "pdf":
            import subprocess, shutil
            pdf_p = REPORTS_DIR / f"{safe}.pdf"

            def _find_bin(*names):
                """Search PATH + common Homebrew/system locations for a binary."""
                extra_dirs = [
                    "/opt/homebrew/bin", "/usr/local/bin",
                    "/usr/bin", "/bin",
                    os.path.expanduser("~/.local/bin"),
                    os.path.expanduser("~/.homebrew/bin"),
                    os.path.expanduser("~/homebrew/bin"),
                    "/Library/TeX/texbin",
                    "/usr/local/texlive/2024/bin/universal-darwin",
                    "/usr/local/texlive/2023/bin/universal-darwin",
                ]
                for nm in names:
                    # 1. Standard PATH lookup
                    found = shutil.which(nm)
                    if found:
                        print(f"[INFO] Found {nm} via PATH: {found}")
                        return found
                    # 2. Explicit directory scan
                    for d in extra_dirs:
                        p = os.path.join(d, nm)
                        if os.path.isfile(p) and os.access(p, os.X_OK):
                            print(f"[INFO] Found {nm} at: {p}")
                            return p
                print(f"[WARN] Binary not found: {names}")
                return None

            # ── Strategy 1: DOCX → PDF via LibreOffice (highest fidelity) ──
            lo = _find_bin("soffice", "libreoffice",
                           "/Applications/LibreOffice.app/Contents/MacOS/soffice")
            if lo:
                try:
                    docx_tmp = REPORTS_DIR / f"{safe}_tmp.docx"
                    markdown_to_docx(content, str(docx_tmp), meta=report_meta)
                    env = {**os.environ, "HOME": str(pathlib.Path.home()), "DISPLAY": ""}
                    r2 = subprocess.run(
                        [lo, "--headless", "--norestore", "--convert-to", "pdf",
                         "--outdir", str(REPORTS_DIR), str(docx_tmp)],
                        capture_output=True, timeout=180, env=env
                    )
                    print(f"[INFO] LibreOffice exit={r2.returncode}")
                    if r2.returncode != 0:
                        print(f"[WARN] LO stderr: {r2.stderr.decode('utf-8','replace')[:400]}")
                    lo_out = REPORTS_DIR / (docx_tmp.stem + ".pdf")
                    if lo_out.exists() and lo_out.stat().st_size > 5000:
                        lo_out.rename(pdf_p)
                        try: docx_tmp.unlink()
                        except: pass
                        print(f"[INFO] PDF via LibreOffice: {pdf_p}")
                        return send_file(str(pdf_p), as_attachment=True, download_name=f"{safe}.pdf")
                    else:
                        print(f"[WARN] LibreOffice did not produce PDF (size={lo_out.stat().st_size if lo_out.exists() else 0})")
                except Exception as e:
                    print(f"[WARN] LibreOffice PDF failed: {e}")

            # ── Strategy 2: pandoc → PDF (wkhtmltopdf engine) ──
            pandoc   = _find_bin("pandoc")
            wk       = _find_bin("wkhtmltopdf")
            if pandoc and wk:
                try:
                    md_p = REPORTS_DIR / f"{safe}_tmp.md"
                    md_p.write_text(content, encoding="utf-8")
                    env2 = {**os.environ, "QT_QPA_PLATFORM": "offscreen",
                            "DISPLAY": ""}
                    r3 = subprocess.run(
                        [pandoc, str(md_p), "-o", str(pdf_p),
                         "--pdf-engine", wk,
                         "--metadata", f"title:{r.get('title','Security Report')}",
                         "-V", "geometry:margin=25mm",
                         "-V", "fontsize=11pt"],
                        capture_output=True, timeout=120, env=env2
                    )
                    if pdf_p.exists() and pdf_p.stat().st_size > 5000:
                        try: md_p.unlink()
                        except: pass
                        print(f"[INFO] PDF via pandoc+wkhtmltopdf: {pdf_p}")
                        return send_file(str(pdf_p), as_attachment=True, download_name=f"{safe}.pdf")
                    else:
                        print(f"[WARN] pandoc+wk failed: {r3.stderr.decode('utf-8','replace')[:400]}")
                except Exception as e:
                    print(f"[WARN] pandoc+wkhtmltopdf failed: {e}")

            # ── Strategy 3: pandoc → PDF (LaTeX / xelatex / pdflatex) ──
            if pandoc:
                xelatex = _find_bin("xelatex", "pdflatex", "lualatex")
                try:
                    md_p = REPORTS_DIR / f"{safe}_tmp.md"
                    md_p.write_text(content, encoding="utf-8")
                    cmd = [pandoc, str(md_p), "-o", str(pdf_p),
                           "--metadata", f"title:{r.get('title','Security Report')}",
                           "-V", "geometry:margin=25mm"]
                    if xelatex:
                        cmd += ["--pdf-engine", xelatex]
                    r4 = subprocess.run(cmd, capture_output=True, timeout=120)
                    if pdf_p.exists() and pdf_p.stat().st_size > 5000:
                        try: md_p.unlink()
                        except: pass
                        print(f"[INFO] PDF via pandoc+LaTeX: {pdf_p}")
                        return send_file(str(pdf_p), as_attachment=True, download_name=f"{safe}.pdf")
                    else:
                        print(f"[WARN] pandoc LaTeX failed: {r4.stderr.decode('utf-8','replace')[:400]}")
                except Exception as e:
                    print(f"[WARN] pandoc LaTeX PDF failed: {e}")

            # ── Strategy 4: pandoc → HTML → PDF via weasyprint ──
            if pandoc:
                weasyprint = _find_bin("weasyprint")
                if weasyprint:
                    try:
                        md_p  = REPORTS_DIR / f"{safe}_tmp.md"
                        html_p= REPORTS_DIR / f"{safe}_tmp.html"
                        md_p.write_text(content, encoding="utf-8")
                        r5a = subprocess.run(
                            [pandoc, str(md_p), "-o", str(html_p), "--standalone"],
                            capture_output=True, timeout=60)
                        if html_p.exists():
                            r5b = subprocess.run(
                                [weasyprint, str(html_p), str(pdf_p)],
                                capture_output=True, timeout=120)
                            if pdf_p.exists() and pdf_p.stat().st_size > 5000:
                                try: md_p.unlink(); html_p.unlink()
                                except: pass
                                print(f"[INFO] PDF via pandoc+weasyprint: {pdf_p}")
                                return send_file(str(pdf_p), as_attachment=True, download_name=f"{safe}.pdf")
                    except Exception as e:
                        print(f"[WARN] weasyprint PDF failed: {e}")

            # ── Strategy 5: Python weasyprint (no pandoc needed) ──
            try:
                import weasyprint as _wp
                md_p = REPORTS_DIR / f"{safe}_tmp.md"
                md_p.write_text(content, encoding="utf-8")
                # Convert markdown to simple HTML first
                try:
                    import markdown as _md_lib
                    html_body = _md_lib.markdown(content, extensions=["tables","fenced_code"])
                except ImportError:
                    # Fallback: minimal markdown → html
                    html_body = "<pre>" + content.replace("<","&lt;").replace(">","&gt;") + "</pre>"
                html_full = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
body{{font-family:Calibri,Arial,sans-serif;margin:25mm;font-size:11pt;color:#222;}}
h1{{color:#1B2A4A;border-bottom:2px solid #1B2A4A;}}
h2{{color:#1F5C99;}} h3{{color:#2E75B6;}}
table{{border-collapse:collapse;width:100%;margin:12px 0;}}
th{{background:#1B2A4A;color:#fff;padding:8px;text-align:left;}}
td{{padding:6px 8px;border:1px solid #C8D4E0;}}
tr:nth-child(even){{background:#F7F9FC;}}
code,pre{{background:#F0F5FA;padding:4px 8px;font-family:Courier New,monospace;font-size:10pt;}}
</style></head><body>{html_body}</body></html>"""
                _wp.HTML(string=html_full).write_pdf(str(pdf_p))
                if pdf_p.exists() and pdf_p.stat().st_size > 5000:
                    try: md_p.unlink()
                    except: pass
                    print(f"[INFO] PDF via Python weasyprint: {pdf_p}")
                    return send_file(str(pdf_p), as_attachment=True, download_name=f"{safe}.pdf")
            except ImportError:
                print("[WARN] weasyprint not installed (pip3 install weasyprint)")
            except Exception as e:
                print(f"[WARN] Python weasyprint failed: {e}")

            # ── All strategies failed ──
            checked = []
            for nm in ["soffice","libreoffice","pandoc","wkhtmltopdf"]:
                b = _find_bin(nm)
                checked.append(f"  {nm}: {'✓ '+b if b else '✗ not found'}")
            try:
                import weasyprint
                checked.append("  weasyprint (Python): ✓ installed")
            except ImportError:
                checked.append("  weasyprint (Python): ✗ not installed")
            return jsonify({"error": (
                "PDF export failed. None of the PDF backends succeeded.\n\n"
                "Binaries / packages checked:\n" + "\n".join(checked) + "\n\n"
                "RECOMMENDED — install weasyprint (no extra dependencies):\n"
                "  pip3 install weasyprint markdown --break-system-packages\n\n"
                "OR install LibreOffice (best quality, converts your DOCX):\n"
                "  macOS:  brew install --cask libreoffice\n"
                "  Linux:  sudo apt install libreoffice\n\n"
                "OR install basictex for pandoc+LaTeX:\n"
                "  macOS:  brew install basictex && eval \"$(/usr/libexec/path_helper)\"\n"
                "  Linux:  sudo apt install pandoc texlive-xetex"
            )}), 500

        return jsonify({"error": "Unknown format"}), 400

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"Export failed: {str(e)}"}), 500


@app.route("/api/report/<int:rid>/shares", methods=["GET"])
@login_required
def get_report_shares(rid):
    """List users this report is shared with."""
    with get_db() as conn:
        r = conn.execute("SELECT created_by,visibility FROM reports WHERE id=?", (rid,)).fetchone()
    if not r:
        return jsonify({"error": "Report not found"}), 404
    r = dict(r)
    if not can_modify_report(r, session.get("username"), session.get("role")):
        return jsonify({"error": "Only the report owner or admin can manage sharing"}), 403
    with get_db() as conn:
        shares = conn.execute(
            "SELECT shared_with, shared_by, created_at FROM report_shares WHERE report_id=? ORDER BY created_at",
            (rid,)
        ).fetchall()
        # Also get list of all users for the share picker
        users = conn.execute(
            "SELECT username, full_name, role FROM users WHERE is_disabled=0 AND username!=? ORDER BY username",
            (session.get("username"),)
        ).fetchall()
    return jsonify({
        "shares": [dict(s) for s in shares],
        "users": [dict(u) for u in users],
        "visibility": r["visibility"],
        "created_by": r["created_by"],
    })


@app.route("/api/report/<int:rid>/visibility", methods=["POST"])
@login_required
def change_visibility(rid):
    """Change report visibility. Owner or admin only."""
    with get_db() as conn:
        r = conn.execute("SELECT created_by,visibility FROM reports WHERE id=?", (rid,)).fetchone()
    if not r:
        return jsonify({"error": "Report not found"}), 404
    r = dict(r)
    if not can_modify_report(r, session.get("username"), session.get("role")):
        return jsonify({"error": "Only the report owner or admin can change visibility"}), 403
    data = request.json or {}
    vis = data.get("visibility", "team")
    if vis not in ("private", "team", "public"):
        return jsonify({"error": "Invalid visibility. Use: private, team, public"}), 400
    with get_db() as conn:
        conn.execute("UPDATE reports SET visibility=? WHERE id=?", (vis, rid))
        conn.commit()
    return jsonify({"success": True, "visibility": vis})


@app.route("/api/report/<int:rid>/share", methods=["POST"])
@login_required
def share_report(rid):
    """Share report with a specific user."""
    with get_db() as conn:
        r = conn.execute("SELECT created_by,visibility FROM reports WHERE id=?", (rid,)).fetchone()
    if not r:
        return jsonify({"error": "Report not found"}), 404
    r = dict(r)
    if not can_modify_report(r, session.get("username"), session.get("role")):
        return jsonify({"error": "Only the report owner or admin can share this report"}), 403
    data = request.json or {}
    share_with = (data.get("username") or "").strip().lower()
    if not share_with:
        return jsonify({"error": "Username required"}), 400
    # Verify target user exists and is not disabled
    with get_db() as conn:
        target = conn.execute(
            "SELECT username FROM users WHERE username=? AND is_disabled=0", (share_with,)
        ).fetchone()
    if not target:
        return jsonify({"error": f"User '{share_with}' not found or is disabled"}), 404
    if share_with == session.get("username"):
        return jsonify({"error": "Cannot share with yourself"}), 400
    try:
        with get_db() as conn:
            conn.execute(
                "INSERT OR IGNORE INTO report_shares (report_id, shared_with, shared_by) VALUES (?,?,?)",
                (rid, share_with, session.get("username"))
            )
            conn.commit()
        return jsonify({"success": True, "shared_with": share_with})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/report/<int:rid>/share/<username>", methods=["DELETE"])
@login_required
def revoke_share(rid, username):
    """Revoke a user's access to a report."""
    with get_db() as conn:
        r = conn.execute("SELECT created_by FROM reports WHERE id=?", (rid,)).fetchone()
    if not r:
        return jsonify({"error": "Report not found"}), 404
    r = dict(r)
    if not can_modify_report(r, session.get("username"), session.get("role")):
        return jsonify({"error": "Only the report owner or admin can revoke access"}), 403
    with get_db() as conn:
        conn.execute(
            "DELETE FROM report_shares WHERE report_id=? AND shared_with=?",
            (rid, username)
        )
        conn.commit()
    return jsonify({"success": True})


@app.route("/api/report/<int:rid>", methods=["DELETE"])
@login_required
def delete_report(rid):
    with get_db() as conn:
        r = conn.execute("SELECT created_by FROM reports WHERE id=?", (rid,)).fetchone()
    if not r:
        return jsonify({"error": "Report not found"}), 404
    # Only the creator or an admin can delete
    if not can_modify_report(r, session.get("username"), session.get("role", "viewer")):
        return jsonify({"error": "You can only delete your own reports"}), 403
    with get_db() as conn:
        conn.execute("DELETE FROM reports WHERE id=?", (rid,))
        conn.commit()
    return jsonify({"success": True})


if __name__ == "__main__":
    print("\n" + "=" * 55)
    print("  VAJRA — Vulnerability Analysis, Judgement & Reporting Arsenal")
    print("=" * 55)
    api_key  = os.environ.get("AI_API_KEY", "")
    base_url = os.environ.get("AI_BASE_URL", "")
    model    = os.environ.get("AI_MODEL", "claude-sonnet-4-6")
    provider = os.environ.get("AI_PROVIDER", "anthropic")
    print(f"  URL:       http://localhost:5000")
    print(f"  Provider:  {provider}")
    print(f"  API Key:   {'SET ✓' if api_key else '⚠ NOT SET'}")
    print(f"  Endpoint:  {base_url if base_url else '⚠ NOT SET — add AI_BASE_URL to .env'}")
    print(f"  Model:     {model}")
    print(f"  Reports:   {REPORTS_DIR}")
    print(f"  Knowledge: {KNOWLEDGE_DIR}")
    print("=" * 55)
    if not api_key or not base_url:
        print("\n  ⚠  Edit .env and set AI_API_KEY, AI_BASE_URL, and AI_PROVIDER\n")
    if not os.environ.get("SECRET_KEY"):
        print("  [WARN] SECRET_KEY not set in .env — sessions will reset on restart!")
        print("  [WARN] Add: SECRET_KEY=your-random-32-char-string\n")
    host = os.environ.get("HOST", "0.0.0.0")
    port = int(os.environ.get("PORT", 5000))
    try:
        from waitress import serve
        print(f"  Server: waitress (production)")
        print(f"  Access: http://<your-server-ip>:{port}")
        serve(app, host=host, port=port, threads=8)
    except ImportError:
        print(f"  Server: Flask dev (install waitress for production)")
        print(f"  Access: http://<your-server-ip>:{port}")
        app.run(host=host, port=port, debug=False)


# ── Manual Report Generation (no AI) ──────────────────────────────────────────

@app.route("/api/generate/manual", methods=["POST"])
@login_required
def generate_manual():
    """
    Generate a DOCX report without AI — structures the user's findings
    directly into the report template using a clean markdown scaffold.
    """
    if session.get("role") == "viewer":
        return jsonify({"error": "Viewer role cannot generate reports."}), 403

    data = request.json or {}
    report_type     = data.get("report_type", "pentest")
    client          = data.get("client", "")
    tester          = data.get("tester", "")
    date            = data.get("date", datetime.now().strftime("%Y-%m-%d"))
    scope           = data.get("scope", "")
    target_system   = data.get("target_system", "")
    tested_from     = data.get("tested_from", "")
    assessment_team = data.get("assessment_team", "")
    engagement_ref  = data.get("engagement_ref", "") or f"VAJRA-{datetime.now().year}-{str(int(datetime.now().timestamp()))[-4:]}"
    classification  = data.get("classification", "CONFIDENTIAL — NOT FOR DISTRIBUTION")
    report_version  = data.get("report_version", "1.0 — Final")
    template_name   = data.get("template_name", "")
    findings_text   = data.get("findings_text", "")
    extra_fields    = data.get("extra_fields", {}) or {}
    title           = data.get("title", f"{report_type.upper()} — {client or 'Client'} — {date}")

    if not findings_text.strip():
        return jsonify({"error": "No findings provided. Please add at least one finding."}), 400

    # ── Build structured markdown scaffold ──────────────────────────────────
    lines = []
    rt_label = {"pentest": "Penetration Test", "redteam": "Red Team", "phishing": "Phishing Campaign"}.get(report_type, report_type.title())

    lines.append(f"# {rt_label} Report — {client or 'Client'}")
    lines.append(f"\n**Engagement Reference:** {engagement_ref}")
    lines.append(f"**Classification:** {classification}")
    lines.append(f"**Report Version:** {report_version}")
    lines.append(f"**Date:** {date}")
    lines.append(f"**Lead Assessor:** {tester}")
    if assessment_team:
        lines.append(f"**Assessment Team:** {assessment_team}")
    lines.append("")

    # Executive Summary placeholder
    lines.append("## Executive Summary")
    lines.append(f"This report presents the findings of the {rt_label.lower()} conducted against **{client or 'the client organisation'}**.")
    if scope:
        lines.append(f"The assessment covered: {scope}.")
    if target_system:
        lines.append(f"Target system: {target_system}.")
    if tested_from:
        lines.append(f"Testing conducted from: {tested_from}.")
    lines.append("")

    # Type-specific metadata
    if report_type == "pentest":
        out_of_scope = extra_fields.get("out_of_scope", "")
        assess_type  = extra_fields.get("assess_type", "")
        if assess_type:
            lines.append(f"**Assessment Type:** {assess_type}")
        if out_of_scope:
            lines.append(f"**Out of Scope:** {out_of_scope}")
        lines.append("")

    elif report_type == "redteam":
        lines.append("## Engagement Overview")
        for key, label in [("objective","Objective"),("threat_actor","Threat Actor Profile"),
                           ("c2_framework","C2 Framework"),("init_access","Initial Access Vector"),
                           ("crown_reached","Crown Jewels Reached"),("detection_rate","Detection Rate"),
                           ("blue_awareness","Blue Team Awareness"),("duration","Duration")]:
            val = extra_fields.get(key, "")
            if val:
                lines.append(f"**{label}:** {val}")
        lines.append("")

    elif report_type == "phishing":
        lines.append("## Campaign Overview")
        for key, label in [("total_recipients","Total Recipients"),("clicked","Links Clicked"),
                           ("creds","Credentials Harvested"),("reported","Reported by Staff"),
                           ("campaign_period","Campaign Period"),("platform","Platform")]:
            val = extra_fields.get(key, "")
            if val:
                lines.append(f"**{label}:** {val}")
        lines.append("")

    # ── Findings ──────────────────────────────────────────────────────────────
    lines.append("## Findings")
    lines.append("")

    parsed = _parse_findings_input(findings_text)

    if parsed:
        for i, f in enumerate(parsed, 1):
            title_f  = f.get("title", f"Finding {i}")
            severity = f.get("severity", "").strip()
            desc     = f.get("description", "").strip()
            impact   = f.get("impact", "").strip()
            remediation = f.get("remediation", "").strip()
            cvss     = f.get("cvss", "").strip()
            ref      = f.get("references", "").strip()

            lines.append(f"### Finding {i}: {title_f}")
            if severity:
                lines.append(f"**Severity:** {severity.upper()}")
            if cvss:
                lines.append(f"**CVSS Score:** {cvss}")
            lines.append("")
            if desc:
                lines.append(f"**Description:**\n{desc}")
                lines.append("")
            if impact:
                lines.append(f"**Impact:**\n{impact}")
                lines.append("")
            if remediation:
                lines.append(f"**Remediation:**\n{remediation}")
                lines.append("")
            if ref:
                lines.append(f"**References:** {ref}")
            lines.append("---")
            lines.append("")
    else:
        # Raw text — include as-is under findings
        lines.append(findings_text.strip())
        lines.append("")

    # ── Conclusion ────────────────────────────────────────────────────────────
    lines.append("## Conclusion")
    lines.append(f"This report was generated manually via VAJRA v1.4 on {date}.")
    lines.append("Please review all findings with your team and prioritise remediation based on severity ratings.")

    md_content = "\n".join(lines)

    # ── Build DOCX ────────────────────────────────────────────────────────────
    import uuid, tempfile
    _tmp_name  = f"manual_tmp_{str(uuid.uuid4())[:8]}"
    out_path   = REPORTS_DIR / f"{_tmp_name}.docx"

    template_path = None
    if template_name:
        tp = TEMPLATES_DIR / template_name
        if tp.exists():
            template_path = str(tp)

    meta = {
        "client":          client,
        "tester":          tester,
        "date":            date,
        "scope":           scope,
        "target_system":   target_system,
        "tested_from":     tested_from,
        "assessment_team": assessment_team,
        "engagement_ref":  engagement_ref,
        "classification":  classification,
        "report_version":  report_version,
        "report_type":     report_type,
        "title":           title,
    }

    try:
        markdown_to_docx(md_content, str(out_path), template_path=template_path, title=title, meta=meta)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"DOCX generation failed: {str(e)}"}), 500

    if not out_path.exists():
        return jsonify({"error": "DOCX file was not created. Check server logs."}), 500
    # DB insert happens after file creation

    # ── Save to DB ────────────────────────────────────────────────────────────
    created_by = session.get("username", "unknown")
    visibility = data.get("visibility", "team")

    with get_db() as conn:
        # Use same schema as AI generate route — report_type not type, no id (AUTOINCREMENT)
        cur = conn.execute("""
            INSERT INTO reports (title, report_type, client, tester, scope,
                target_system, tested_from, assessment_team, engagement_ref,
                classification, report_version, status, created_by, visibility,
                findings_raw, generated_content, template_used)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,'final',?,?,?,?,?)
        """, (title, report_type, client, tester, scope,
              target_system, tested_from, assessment_team, engagement_ref,
              classification, report_version, created_by, visibility,
              findings_text, md_content, template_name or ''))
        db_id = cur.lastrowid
        conn.commit()

    # Rename file to use db id for consistency with download route
    final_path = REPORTS_DIR / f"{db_id}.docx"
    import shutil
    shutil.move(str(out_path), str(final_path))

    return jsonify({
        "report_id":   db_id,
        "download_url": f"/api/export/{db_id}/docx",
        "manual":      True,
    })
